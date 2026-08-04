from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path


PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "tools"))

import kb_raw_structure as structure
import kb_search


class RawStructureGoldenTests(unittest.TestCase):
    def test_html_preserves_block_and_table_order(self) -> None:
        source = PROJECT_ROOT / "tests" / "fixtures" / "raw_structure" / "official.html"
        result = structure.extract_html(source, "Fixture notice", "faithful")

        self.assertLess(result.body.index("Opening paragraph."), result.body.index("## Section one"))
        self.assertLess(result.body.index("Paragraph before the table."), result.body.index("| Item | Amount |"))
        self.assertLess(result.body.index("| Item | Amount |"), result.body.index("Paragraph after the table."))
        self.assertNotIn("Navigation noise", result.body)
        self.assertNotIn("Footer noise", result.body)
        self.assertEqual(result.tables[0].representation, "markdown")

    def test_html_uses_tolerant_parser_for_malformed_official_snapshot(self) -> None:
        source_text = "<html><head><title>Broken/title><body><div id='UCAP-CONTENT'><p>" + ("Official body. " * 30) + "</p></div></body></html>"
        with tempfile.TemporaryDirectory() as temporary:
            source = Path(temporary) / "malformed.html"
            source.write_text(source_text, encoding="utf-8")
            result = structure.extract_html(source, "Malformed fixture", "faithful")
        self.assertIn("Official body.", result.body)
        self.assertIn("#UCAP-CONTENT", result.engine)

    def test_docx_keeps_paragraph_table_paragraph_sequence(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "fixture.docx"
            document = Document()
            document.add_heading("Section A", level=1)
            document.add_paragraph("Before table.")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Key"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "A"
            table.cell(1, 1).text = "1"
            document.add_paragraph("After table.")
            document.save(path)

            result = structure.extract_docx(path, "DOCX fixture", "faithful")

        self.assertIn("## Section A", result.body)
        self.assertLess(result.body.index("Before table."), result.body.index("| Key | Value |"))
        self.assertLess(result.body.index("| Key | Value |"), result.body.index("After table."))
        self.assertEqual(result.engine, "python-docx-xml-order")

    def test_pdf_adds_page_markers_and_removes_repeated_header(self) -> None:
        import fitz

        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "fixture.pdf"
            document = fitz.open()
            for page_number in range(1, 4):
                page = document.new_page()
                page.insert_text((72, 35), "Repeated official header", fontsize=9)
                page.insert_text((72, 100), f"Page {page_number} section", fontsize=16)
                page.insert_text((72, 135), f"Body text for page {page_number}.", fontsize=11)
                page.insert_text((280, 810), str(page_number), fontsize=9)
            document.save(path)
            document.close()

            result = structure.extract_pdf(path, "PDF fixture", "faithful")

        self.assertEqual(result.source_pages, 3)
        self.assertEqual(structure.PAGE_MARKER_RE.findall(result.body), ["1", "2", "3"])
        self.assertNotIn("Repeated official header", result.body)
        self.assertIn("Body text for page 3.", result.body)

    def test_structure_json_is_search_excluded(self) -> None:
        self.assertTrue(kb_search.is_search_excluded("raw/laws/example.structure.json"))
        self.assertFalse(kb_search.is_search_excluded("raw/laws/example.json"))

    def test_markdown_passthrough_preserves_body(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "lecture.md"
            path.write_text("# Lesson\n\n```markdown\n# Example heading\n```\n", encoding="utf-8")
            result = structure.extract_markdown_passthrough(path)
        self.assertEqual(result.body, "# Lesson\n\n```markdown\n# Example heading\n```\n")
        self.assertEqual(result.engine, "markdown-pass-through")
        self.assertEqual(len(result.headings), 1)

    def test_frontmatter_update_is_idempotent(self) -> None:
        first = structure.update_frontmatter("title: Example", {"source_pages": 2, "representation": "semantic-transcript"})
        match = structure.FRONTMATTER_RE.match(first)
        self.assertIsNotNone(match)
        second = structure.update_frontmatter(match.group(1), {"source_pages": 2, "representation": "semantic-transcript"})
        self.assertEqual(first, second)

    def test_heading_levels_are_clamped_without_false_duplicate_sections(self) -> None:
        rendered = ["# Document", "## Chapter"]
        headings = [
            structure.Heading(1, "Document", None, "test", "test"),
            structure.Heading(2, "Chapter", None, "test", "test"),
        ]
        level = structure.append_heading(rendered, headings, 4, "Article", None, "test", "test")
        self.assertEqual(level, 3)
        self.assertEqual(structure.heading_issues("# Doc\n\n## General\n\n### Detail\n\n## General\n"), [])
        self.assertTrue(structure.heading_issues("# Doc\n\n# Doc\n"))

    def test_audit_writes_only_requested_output_by_default(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            output = root / "reports" / "audit.md"
            result = structure.audit(root, "all", output)
            self.assertEqual(0, result)
            self.assertTrue(output.exists())
            self.assertFalse((root / "wiki" / "_maintenance" / "raw-structure-review.md").exists())


if __name__ == "__main__":
    unittest.main()
