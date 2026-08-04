from __future__ import annotations

import hashlib
import re
import sys
from datetime import date
from pathlib import Path

import docx  # type: ignore

import kb_pdf_to_markdown


ROOT = Path("knowledge-base/CPA-ZH")
TARGET_DIR = ROOT / "raw" / "standards" / "accounting" / "interpretations-pages"


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def rel(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()


def yaml_quote(value: str) -> str:
    escaped = value.replace("\\", "\\\\").replace('"', '\\"')
    return f'"{escaped}"'


def body_without_frontmatter(text: str) -> str:
    return re.sub(r"^---\n.*?\n---\n?", "", text, flags=re.S).strip()


def markdown_from_text(source: Path, title: str, body: str, method: str, note: str = "") -> str:
    lines = [
        "---",
        f"title: {yaml_quote(title)}",
        "type: extracted-text",
        "source_type: interpretation-attachment-markdown",
        f"created: {date.today().isoformat()}",
        f"updated: {date.today().isoformat()}",
        f"original_file: {yaml_quote(rel(source))}",
        f"source_sha256: {sha256_file(source)}",
        f"extraction_method: {yaml_quote(method)}",
        f"text_length: {len(body)}",
    ]
    if note:
        lines.append(f"conversion_note: {yaml_quote(note)}")
    lines.extend(
        [
            "tags: [企业会计准则, 准则解释, markdown派生件]",
            "---",
            "",
            f"# {title}",
            "",
        ]
    )
    if note:
        lines.extend([f"> 转换说明：{note}", ""])
    lines.extend(["## 抽取正文", "", body.strip(), ""])
    return "\n".join(lines)


def docx_to_text(path: Path) -> str:
    document = docx.Document(str(path))
    paragraphs: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs).strip()


def matching_html_markdown(number: int) -> Path | None:
    patterns = [f"第{number}号", f"第{number:02d}号"]
    for path in sorted(TARGET_DIR.glob("*.html.md")):
        name = path.name
        if any(pattern in name for pattern in patterns):
            return path
    return None


def doc_fallback_text(path: Path, number: int) -> tuple[str, str]:
    html_md = matching_html_markdown(number)
    if not html_md:
        return "", "未找到同号官方 HTML Markdown，需人工转换老 .doc。"
    body = body_without_frontmatter(html_md.read_text(encoding="utf-8", errors="ignore"))
    body = re.sub(r"(?m)^# ", "### ", body)
    note = f"本机无可用 Word/LibreOffice 转换老 .doc；本派生件采用同目录同号官方 HTML Markdown 正文作为检索文本，原 .doc 仍保留归档。来源：{rel(html_md)}"
    return body, note


def convert_one(path: Path, overwrite: bool) -> tuple[str, int]:
    output = Path(str(path) + ".md")
    if output.exists() and not overwrite:
        return "skip-existing", 0

    number_match = re.search(r"第(\d+)号", path.stem)
    number = int(number_match.group(1)) if number_match else 0
    title = path.stem
    suffix = path.suffix.lower()
    note = ""

    if suffix == ".pdf":
        body, engine, quality = kb_pdf_to_markdown.extract_pdf_best(path, preferred="auto")
        method = f"pdf:{engine}; quality={quality}"
    elif suffix == ".docx":
        body = docx_to_text(path)
        method = "python-docx"
    elif suffix == ".doc":
        body, note = doc_fallback_text(path, number)
        method = "official-html-md-fallback"
    else:
        return "unsupported", 0

    if not body.strip():
        return "empty", 0

    markdown = markdown_from_text(path, title, body, method, note)
    output.write_text(markdown, encoding="utf-8", newline="\n")
    return "written", len(body)


def main(argv: list[str]) -> int:
    overwrite = "--overwrite" in argv
    if not TARGET_DIR.exists():
        print(f"missing target directory: {TARGET_DIR}", file=sys.stderr)
        return 2

    targets = [
        path
        for path in sorted(TARGET_DIR.iterdir())
        if path.is_file()
        and path.suffix.lower() in {".doc", ".docx", ".pdf"}
        and not Path(str(path) + ".md").exists()
    ]
    if overwrite:
        targets = [
            path
            for path in sorted(TARGET_DIR.iterdir())
            if path.is_file() and path.suffix.lower() in {".doc", ".docx", ".pdf"}
        ]

    counts: dict[str, int] = {}
    for path in targets:
        status, length = convert_one(path, overwrite=overwrite)
        counts[status] = counts.get(status, 0) + 1
        print(f"{status}: {rel(path)} chars={length}")

    print("summary:", " ".join(f"{key}={value}" for key, value in sorted(counts.items())))
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
