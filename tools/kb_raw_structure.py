from __future__ import annotations

import argparse
import hashlib
import html
import json
import re
import shutil
import statistics
import sys
from collections import Counter
from dataclasses import dataclass, field
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any, Iterable

import yaml


FRONTMATTER_RE = re.compile(r"\A---\s*\n(.*?)\n---\s*\n?", re.DOTALL)
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$", re.MULTILINE)
PAGE_MARKER_RE = re.compile(r"^<!-- source-page: (\d+) -->$", re.MULTILINE)
LEGAL_TOKEN_RE = re.compile(
    r"^第[零〇一二三四五六七八九十百千万亿两0-9]+(?P<kind>编|章|节|条)(?:\s|　|[:：]|$)"
)
QUESTION_RE = re.compile(r"^(?:问题|问)\s*[：:]\s*")
ANSWER_RE = re.compile(r"^(?:答复|回答|答)\s*[：:]\s*")
SUPPORTED_SUFFIXES = {".html", ".htm", ".pdf", ".docx", ".doc", ".md"}
CONTENT_ROLES = {"content", "substantive-attachment", "equivalent-source"}


@dataclass
class Heading:
    level: int
    text: str
    page: int | None
    source: str
    locator: str
    confidence: float = 1.0


@dataclass
class TableInfo:
    index: int
    page: int | None
    rows: int
    columns: int
    representation: str
    locator: str
    needs_review: bool = False


@dataclass
class ExtractionResult:
    body: str
    engine: str
    source_pages: int
    headings: list[Heading] = field(default_factory=list)
    tables: list[TableInfo] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    review_pages: list[int] = field(default_factory=list)
    source_text_chars: int = 0
    extracted_text_chars: int = 0

    @property
    def coverage(self) -> float:
        if self.source_text_chars <= 0:
            return 1.0 if self.extracted_text_chars else 0.0
        return min(1.0, self.extracted_text_chars / self.source_text_chars)

    @property
    def status(self) -> str:
        return "needs-review" if self.warnings or self.review_pages else "verified-auto"


@dataclass
class Target:
    markdown: Path
    original: Path | None
    metadata: dict[str, Any]
    category: str


def utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def read_text(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def split_markdown(path: Path) -> tuple[dict[str, Any], str, str]:
    text = read_text(path)
    match = FRONTMATTER_RE.match(text)
    if not match:
        return {}, "", text
    try:
        metadata = yaml.safe_load(match.group(1)) or {}
    except yaml.YAMLError:
        metadata = {}
    return metadata if isinstance(metadata, dict) else {}, match.group(1), text[match.end() :]


def yaml_scalar(value: Any) -> str:
    if value is None:
        return "null"
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, (int, float)):
        return str(value)
    return json.dumps(str(value), ensure_ascii=False)


def update_frontmatter(raw_frontmatter: str, updates: dict[str, Any]) -> str:
    lines = raw_frontmatter.splitlines() if raw_frontmatter else []
    for key, value in updates.items():
        replacement = f"{key}: {yaml_scalar(value)}"
        pattern = re.compile(rf"^{re.escape(key)}\s*:")
        for index, line in enumerate(lines):
            if pattern.match(line):
                lines[index] = replacement
                break
        else:
            lines.append(replacement)
    return "---\n" + "\n".join(lines).strip() + "\n---\n\n"


def clean_inline(text: str) -> str:
    text = text.replace("\u00a0", " ").replace("\u3000", " ").replace("\ufeff", "")
    return re.sub(r"[ \t]+", " ", text).strip()


def comparable_text(text: str) -> str:
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"[#*_`|>\-]", "", text)
    return re.sub(r"\s+", "", html.unescape(text))


def markdown_heading_records(body: str) -> list[tuple[int, int, str]]:
    records: list[tuple[int, int, str]] = []
    fence: str | None = None
    for line_number, line in enumerate(body.splitlines(), 1):
        fence_match = re.match(r"^\s*(```+|~~~+)", line)
        if fence_match:
            marker = fence_match.group(1)[0]
            fence = None if fence == marker else marker if fence is None else fence
            continue
        if fence is not None:
            continue
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if match:
            records.append((line_number, len(match.group(1)), clean_inline(match.group(2))))
    return records


def markdown_title(metadata: dict[str, Any], path: Path) -> str:
    return clean_inline(str(metadata.get("title") or path.name.removesuffix(".md")))


def relative(root: Path, path: Path) -> str:
    try:
        return path.relative_to(root).as_posix()
    except ValueError:
        return path.as_posix()


def classify_path(root: Path, markdown: Path) -> str:
    rel = relative(root, markdown)
    if rel.startswith("raw/lectures/"):
        return "lectures"
    if rel.startswith("raw/cases/") or any(
        part in rel
        for part in (
            "/application-case-attachments/",
            "/application-cases-pages/",
            "/implementation-qa-pages-v2/",
        )
    ):
        return "cases"
    if rel.startswith(("raw/laws/", "raw/standards/", "raw/policies/", "raw/ethics/")):
        return "authoritative"
    return "other"


def resolve_original(root: Path, raw_value: str, markdown: Path) -> Path | None:
    candidates: list[Path] = []
    if raw_value:
        supplied = Path(raw_value.replace("\\", "/"))
        candidates.append(supplied if supplied.is_absolute() else root / supplied)
        parts = list(supplied.parts)
        if parts and parts[0].lower() == "raw" and (len(parts) < 2 or parts[1] != "_archive"):
            candidates.append(root / "raw" / "_archive" / Path(*parts[1:]))
    candidates.append(Path(str(markdown)[: -len(".md")]))
    for candidate in candidates:
        if candidate.exists() and candidate.is_file():
            return candidate.resolve()

    if raw_value:
        wanted = Path(raw_value.replace("\\", "/"))
        matches = list((root / "raw" / "_archive").rglob(wanted.name))
        if len(matches) == 1:
            return matches[0].resolve()
        if matches:
            wanted_parts = tuple(part.lower() for part in wanted.parts[-4:])
            ranked = sorted(
                matches,
                key=lambda path: sum(
                    1 for left, right in zip(reversed(path.parts), reversed(wanted_parts)) if left.lower() == right
                ),
                reverse=True,
            )
            return ranked[0].resolve()
    return None


def iter_targets(root: Path, scope: str) -> Iterable[Target]:
    for markdown in sorted((root / "raw").rglob("*.md")):
        if "_archive" in markdown.parts or markdown.name == "README.md":
            continue
        metadata, _, _ = split_markdown(markdown)
        category = classify_path(root, markdown)
        if scope != "all" and category != scope:
            continue
        if category == "other":
            continue
        role = str(metadata.get("source_role") or "content")
        if role not in CONTENT_ROLES:
            continue
        raw_value = str(metadata.get("original_file") or "")
        original = resolve_original(root, raw_value, markdown)
        if original is None and not raw_value and category in {"authoritative", "lectures"}:
            original = markdown.resolve()
        yield Target(markdown, original, metadata, category)


def legal_levels(texts: Iterable[str]) -> dict[str, int]:
    kinds = {match.group("kind") for text in texts if (match := LEGAL_TOKEN_RE.match(text))}
    order = [kind for kind in ("编", "章", "节", "条") if kind in kinds]
    return {kind: min(6, index + 2) for index, kind in enumerate(order)}


def render_markdown_table(rows: list[list[str]]) -> str:
    width = max((len(row) for row in rows), default=0)
    if not width:
        return ""
    normalized = [[clean_inline(cell).replace("|", "\\|") for cell in row] + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    body = normalized[1:]
    return "\n".join(
        ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * width) + " |"]
        + ["| " + " | ".join(row) + " |" for row in body]
    )


def render_html_table(rows: list[list[str]], review: bool = False) -> str:
    lines = []
    if review:
        lines.append("<!-- table-layout: complex; verify against the source -->")
    lines.append("<table>")
    for row in rows:
        lines.append("  <tr>")
        for cell in row:
            lines.append(f"    <td>{html.escape(clean_inline(cell))}</td>")
        lines.append("  </tr>")
    lines.append("</table>")
    return "\n".join(lines)


def append_heading(
    rendered: list[str],
    headings: list[Heading],
    desired_level: int,
    text: str,
    page: int | None,
    source: str,
    locator: str,
    confidence: float = 1.0,
) -> int:
    previous = headings[-1].level if headings else 0
    level = min(max(1, desired_level), min(6, previous + 1)) if previous else min(6, max(1, desired_level))
    rendered.append(f"{'#' * level} {text}")
    headings.append(Heading(level, text, page, source, locator, confidence))
    return level


def table_is_simple(rows: list[list[str]], has_merge: bool = False) -> bool:
    if has_merge or not rows:
        return False
    widths = {len(row) for row in rows}
    return len(widths) == 1 and next(iter(widths), 0) > 0 and all("\n" not in cell for row in rows for cell in row)


def choose_html_container(soup: Any) -> tuple[Any, str, bool]:
    selectors = (
        "#UCAP-CONTENT",
        "#zoom",
        ".TRS_Editor",
        ".pages_content",
        ".trs_paper_default",
        ".article-content",
        ".article_content",
        ".detail-content",
        ".content-body",
        ".content",
        "article",
        "main",
    )
    for selector in selectors:
        candidates: list[tuple[int, Any, str]] = []
        for node in soup.select(selector):
            score = len(clean_inline(node.get_text(" ", strip=True))) + 250 * len(node.find_all(["p", "table"]))
            candidates.append((score, node, selector))
        if candidates:
            score, node, selected = max(candidates, key=lambda item: item[0])
            if score >= 100:
                return node, selected, False
    return soup.body or soup, "body-fallback", True


def extract_html(path: Path, title: str, profile: str) -> ExtractionResult:
    from bs4 import BeautifulSoup, NavigableString, Tag

    source = read_text(path)
    parsed = [BeautifulSoup(source, parser) for parser in ("lxml", "html.parser")]
    soup = max(parsed, key=lambda item: len(clean_inline(item.get_text(" ", strip=True))))
    for node in soup.select("script,style,noscript,nav,footer,header,form,aside"):
        node.decompose()
    container, selector, fallback = choose_html_container(soup)
    headings: list[Heading] = []
    tables: list[TableInfo] = []
    warnings = ["html-content-container-fallback"] if fallback else []
    blocks: list[tuple[str, int | None, str, str]] = []
    ordinal = Counter()

    legal_candidates = [clean_inline(node.get_text(" ", strip=True)) for node in container.find_all(["p", "div"])]
    levels = legal_levels(legal_candidates) if profile == "faithful" else {}

    def walk(node: Any) -> None:
        if isinstance(node, NavigableString) or not isinstance(node, Tag):
            return
        tag = node.name.lower()
        ordinal[tag] += 1
        locator = f"{selector} {tag}[{ordinal[tag]}]"
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            value = clean_inline(node.get_text(" ", strip=True))
            if value:
                level = int(tag[1])
                blocks.append((value, level, "html-tag", locator))
            return
        if tag == "table":
            rows = [[clean_inline(cell.get_text(" ", strip=True)) for cell in row.find_all(["th", "td"], recursive=False)] for row in node.find_all("tr")]
            rows = [row for row in rows if row]
            has_merge = any(cell.has_attr("rowspan") or cell.has_attr("colspan") for cell in node.find_all(["th", "td"]))
            simple = table_is_simple(rows, has_merge)
            index = len(tables) + 1
            tables.append(TableInfo(index, None, len(rows), max((len(row) for row in rows), default=0), "markdown" if simple else "html", locator, not simple))
            blocks.append((render_markdown_table(rows) if simple else render_html_table(rows, review=has_merge), None, "table", locator))
            if has_merge:
                warnings.append(f"complex-html-table:{index}")
            return
        if tag in {"ul", "ol"}:
            ordered = tag == "ol"
            for index, item in enumerate(node.find_all("li", recursive=False), 1):
                value = clean_inline(item.get_text(" ", strip=True))
                if value:
                    blocks.append((f"{index}. {value}" if ordered else f"- {value}", None, "list", locator))
            return
        if tag in {"p", "pre", "blockquote"}:
            value = clean_inline(node.get_text(" ", strip=True))
            if value:
                match = LEGAL_TOKEN_RE.match(value)
                level = levels.get(match.group("kind")) if match and profile == "faithful" else None
                blocks.append((value, level, "explicit-legal-numbering" if level else "paragraph", locator))
            return
        for child in node.children:
            walk(child)

    walk(container)
    rendered = [f"# {title}"]
    headings.append(Heading(1, title, None, "frontmatter-title", "frontmatter.title"))
    for value, level, source, locator in blocks:
        if not value or (level and clean_inline(value) == title and len(rendered) == 1):
            continue
        if level:
            append_heading(rendered, headings, level, value, None, source, locator)
        else:
            rendered.append(value)
    source_chars = len(comparable_text(container.get_text(" ", strip=True)))
    body = "\n\n".join(rendered).strip() + "\n"
    parser_name = getattr(getattr(soup, "builder", None), "NAME", "unknown")
    return ExtractionResult(body, f"beautifulsoup-{parser_name}:{selector}", 1, headings, tables, warnings, [], source_chars, len(comparable_text(body)))


def iter_docx_blocks(document: Any) -> Iterable[tuple[str, Any]]:
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield "paragraph", Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield "table", Table(child, document)


def extract_docx(path: Path, title: str, profile: str) -> ExtractionResult:
    from docx import Document

    document = Document(str(path))
    raw_blocks = list(iter_docx_blocks(document))
    levels = legal_levels(
        clean_inline(block.text) for kind, block in raw_blocks if kind == "paragraph" and getattr(block, "text", "")
    ) if profile == "faithful" else {}
    rendered = [f"# {title}"]
    headings = [Heading(1, title, None, "frontmatter-title", "frontmatter.title")]
    tables: list[TableInfo] = []
    warnings: list[str] = []
    paragraph_index = 0
    table_index = 0
    source_text: list[str] = []

    for kind, block in raw_blocks:
        if kind == "paragraph":
            paragraph_index += 1
            value = clean_inline(block.text)
            if not value:
                continue
            source_text.append(value)
            style_name = clean_inline(getattr(getattr(block, "style", None), "name", "") or "")
            style_match = re.search(r"(?:Heading|标题)\s*([1-6])", style_name, re.IGNORECASE)
            level = min(6, int(style_match.group(1)) + 1) if style_match else None
            source = "word-heading-style" if level else "paragraph"
            if level is None and profile == "faithful":
                legal_match = LEGAL_TOKEN_RE.match(value)
                if legal_match:
                    level = levels.get(legal_match.group("kind"))
                    source = "explicit-legal-numbering"
            if level:
                level = min(6, max(1, level))
                if value == title and len(rendered) == 1:
                    continue
                append_heading(
                    rendered,
                    headings,
                    level,
                    value,
                    None,
                    source,
                    f"word/paragraph[{paragraph_index}]",
                )
            else:
                rendered.append(value)
        else:
            table_index += 1
            rows = [[clean_inline(cell.text) for cell in row.cells] for row in block.rows]
            source_text.extend(cell for row in rows for cell in row)
            xml = block._tbl.xml
            has_merge = "w:vMerge" in xml or "w:gridSpan" in xml
            simple = table_is_simple(rows, has_merge)
            tables.append(TableInfo(table_index, None, len(rows), max((len(row) for row in rows), default=0), "markdown" if simple else "html", f"word/table[{table_index}]", not simple))
            rendered.append(render_markdown_table(rows) if simple else render_html_table(rows, review=has_merge))
            if has_merge:
                warnings.append(f"complex-docx-table:{table_index}")

    body = "\n\n".join(rendered).strip() + "\n"
    return ExtractionResult(body, "python-docx-xml-order", 0, headings, tables, warnings, [], len(comparable_text(" ".join(source_text))), len(comparable_text(body)))


def extract_markdown_passthrough(path: Path, legacy_doc: bool = False) -> ExtractionResult:
    _, _, body = split_markdown(path)
    body = body.replace("\r\n", "\n").replace("\r", "\n")
    headings: list[Heading] = []
    for line_number, level, text in markdown_heading_records(body):
        headings.append(Heading(level, text, None, "markdown-heading", f"markdown/line[{line_number}]"))
    table_count = len(re.findall(r"^\|.*\|$", body, re.MULTILINE)) // 3
    tables = [
        TableInfo(index, None, 0, 0, "markdown-existing", f"markdown/table[{index}]", legacy_doc)
        for index in range(1, table_count + 1)
    ]
    warnings = ["fallback-existing-derived-legacy-doc"] if legacy_doc else []
    return ExtractionResult(
        body.rstrip() + "\n",
        "fallback-existing-derived-legacy-doc" if legacy_doc else "markdown-pass-through",
        0,
        headings,
        tables,
        warnings,
        [],
        len(comparable_text(body)),
        len(comparable_text(body)),
    )


def join_pdf_lines(lines: list[str]) -> str:
    result = ""
    for line in lines:
        line = clean_inline(line)
        if not line:
            continue
        if not result:
            result = line
            continue
        if result.endswith("-") and re.search(r"[A-Za-z]-$", result) and re.match(r"^[A-Za-z]", line):
            result = result[:-1] + line
        elif re.search(r"[A-Za-z0-9,;:]$", result) and re.match(r"^[A-Za-z0-9(]", line):
            result += " " + line
        else:
            result += line
    return result


def pdf_line_records(page: Any) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    data = page.get_text("dict", sort=True)
    for block_index, block in enumerate(data.get("blocks", []), 1):
        if block.get("type") != 0:
            continue
        for line_index, line in enumerate(block.get("lines", []), 1):
            spans = line.get("spans", [])
            value = clean_inline("".join(str(span.get("text", "")) for span in spans))
            if not value:
                continue
            bbox = tuple(line.get("bbox") or block.get("bbox") or (0, 0, 0, 0))
            sizes = [float(span.get("size") or 0) for span in spans if span.get("text")]
            flags = [int(span.get("flags") or 0) for span in spans]
            records.append(
                {
                    "text": value,
                    "bbox": bbox,
                    "size": max(sizes, default=0.0),
                    "bold": any(flag & 16 for flag in flags),
                    "block": block_index,
                    "line": line_index,
                }
            )
    return records


def repeated_pdf_margins(pages: list[list[dict[str, Any]]], heights: list[float]) -> set[str]:
    counter: Counter[str] = Counter()
    for records, height in zip(pages, heights):
        seen: set[str] = set()
        for record in records:
            y0, y1 = float(record["bbox"][1]), float(record["bbox"][3])
            if y1 <= height * 0.12 or y0 >= height * 0.88:
                key = re.sub(r"\d+", "#", comparable_text(record["text"]))
                if len(key) >= 3:
                    seen.add(key)
        counter.update(seen)
    threshold = max(3, int(len(pages) * 0.6 + 0.5))
    return {key for key, count in counter.items() if count >= threshold}


def point_in_bbox(x: float, y: float, bbox: tuple[float, float, float, float]) -> bool:
    return bbox[0] <= x <= bbox[2] and bbox[1] <= y <= bbox[3]


def extract_pdf_tables(page: Any, page_number: int) -> tuple[list[dict[str, Any]], list[str]]:
    found: list[dict[str, Any]] = []
    warnings: list[str] = []
    try:
        finder = page.find_tables()
        for index, table in enumerate(finder.tables, 1):
            raw_rows = table.extract() or []
            rows = [[clean_inline(str(cell or "")) for cell in row] for row in raw_rows]
            rows = [row for row in rows if any(row)]
            cells = getattr(table, "cells", []) or []
            has_merge = any(cell is None for cell in cells)
            simple = table_is_simple(rows, has_merge)
            found.append(
                {
                    "bbox": tuple(table.bbox),
                    "rows": rows,
                    "simple": simple,
                    "has_merge": has_merge,
                    "locator": f"pdf/page[{page_number}]/table[{index}]",
                }
            )
            if not simple:
                warnings.append(f"complex-pdf-table:page-{page_number}-table-{index}")
    except Exception as exc:
        warnings.append(f"pdf-table-detection-failed:page-{page_number}:{type(exc).__name__}")
    return found, warnings


def extract_pdf(path: Path, title: str, profile: str) -> ExtractionResult:
    import fitz

    document = fitz.open(str(path))
    try:
        page_records = [pdf_line_records(page) for page in document]
        heights = [float(page.rect.height) for page in document]
        repeated = repeated_pdf_margins(page_records, heights)
        all_sizes = [record["size"] for records in page_records for record in records for _ in range(max(1, len(record["text"]))) if record["size"] > 0]
        body_size = statistics.median(all_sizes) if all_sizes else 10.0
        toc_entries = document.get_toc(simple=True) or []
        toc_map: dict[tuple[int, str], int] = {}
        for level, toc_title, page_number, *_ in toc_entries:
            toc_map[(int(page_number), comparable_text(str(toc_title)))] = min(6, int(level) + 1)
        all_texts = [record["text"] for records in page_records for record in records]
        levels = legal_levels(all_texts) if profile == "faithful" else {}

        rendered = [f"# {title}"]
        headings = [Heading(1, title, None, "frontmatter-title", "frontmatter.title")]
        tables: list[TableInfo] = []
        warnings: list[str] = []
        review_pages: list[int] = []
        source_parts: list[str] = []

        for page_index, (page, records, height) in enumerate(zip(document, page_records, heights), 1):
            rendered.append(f"<!-- source-page: {page_index} -->")
            page_tables, table_warnings = extract_pdf_tables(page, page_index) if profile == "faithful" else ([], [])
            warnings.extend(table_warnings)
            events: list[tuple[float, str, Any]] = []
            for table in page_tables:
                events.append((float(table["bbox"][1]), "table", table))

            grouped: dict[int, list[dict[str, Any]]] = {}
            for record in records:
                y0, y1 = float(record["bbox"][1]), float(record["bbox"][3])
                margin_key = re.sub(r"\d+", "#", comparable_text(record["text"]))
                if margin_key in repeated and (y1 <= height * 0.12 or y0 >= height * 0.88):
                    continue
                center_x = (float(record["bbox"][0]) + float(record["bbox"][2])) / 2
                center_y = (y0 + y1) / 2
                if any(point_in_bbox(center_x, center_y, table["bbox"]) for table in page_tables):
                    continue
                grouped.setdefault(int(record["block"]), []).append(record)

            for block_records in grouped.values():
                block_records.sort(key=lambda item: (item["bbox"][1], item["bbox"][0]))
                y0 = min(float(item["bbox"][1]) for item in block_records)
                events.append((y0, "text", block_records))

            for _, event_type, payload in sorted(events, key=lambda item: (item[0], item[1])):
                if event_type == "table":
                    rows = payload["rows"]
                    table_index = len(tables) + 1
                    simple = bool(payload["simple"])
                    tables.append(TableInfo(table_index, page_index, len(rows), max((len(row) for row in rows), default=0), "markdown" if simple else "html", payload["locator"], not simple))
                    rendered.append(render_markdown_table(rows) if simple else render_html_table(rows, review=True))
                    source_parts.extend(cell for row in rows for cell in row)
                    if not simple:
                        review_pages.append(page_index)
                    continue

                paragraph_lines: list[str] = []
                for record in payload:
                    value = record["text"]
                    source_parts.append(value)
                    locator = f"pdf/page[{page_index}]/block[{record['block']}]/line[{record['line']}]"
                    normalized = comparable_text(value)
                    toc_level = next((level for (page_no, toc_text), level in toc_map.items() if page_no == page_index and (normalized == toc_text or normalized.startswith(toc_text))), None)
                    legal_match = LEGAL_TOKEN_RE.match(value) if profile == "faithful" else None
                    legal_level = levels.get(legal_match.group("kind")) if legal_match else None
                    width = float(record["bbox"][2]) - float(record["bbox"][0])
                    page_width = float(page.rect.width)
                    centered = abs((float(record["bbox"][0]) + float(record["bbox"][2])) / 2 - page_width / 2) < page_width * 0.08
                    font_heading = (
                        profile == "faithful"
                        and len(value) <= 80
                        and record["size"] >= body_size + 1.5
                        and (record["bold"] or centered or width < page_width * 0.7)
                    )
                    level = toc_level or legal_level or (2 if font_heading else None)
                    if level:
                        if paragraph_lines:
                            rendered.append(join_pdf_lines(paragraph_lines))
                            paragraph_lines = []
                        if value == title and len(headings) == 1:
                            continue
                        source = "pdf-toc" if toc_level else "explicit-legal-numbering" if legal_level else "pdf-font-bbox"
                        append_heading(
                            rendered,
                            headings,
                            level,
                            value,
                            page_index,
                            source,
                            locator,
                            1.0 if toc_level or legal_level else 0.85,
                        )
                    else:
                        paragraph_lines.append(value)
                if paragraph_lines:
                    rendered.append(join_pdf_lines(paragraph_lines))

            page_source_chars = len(comparable_text("".join(record["text"] for record in records)))
            if page_source_chars < 20:
                review_pages.append(page_index)
                warnings.append(f"low-pdf-text:page-{page_index}")

        body = "\n\n".join(part for part in rendered if part).strip() + "\n"
        return ExtractionResult(
            body,
            "pymupdf-dict+bbox",
            len(document),
            headings,
            tables,
            sorted(set(warnings)),
            sorted(set(review_pages)),
            len(comparable_text("".join(source_parts))),
            len(comparable_text(body)),
        )
    finally:
        document.close()


def extract_target(target: Target, profile: str) -> ExtractionResult:
    if target.original is None:
        raise FileNotFoundError("original source could not be resolved")
    suffix = target.original.suffix.lower()
    title = markdown_title(target.metadata, target.markdown)
    extraction_engine = str(target.metadata.get("extraction_engine") or "")
    note = str(target.metadata.get("conversion_note") or "")
    fallback_match = re.search(r"(raw[/\\][^()]+?\.html)", note)
    if (
        "official-html-fallback" in extraction_engine
        or "fallback-official-html" in extraction_engine
        or fallback_match is not None
    ):
        if fallback_match:
            raw_root = next((parent for parent in target.markdown.parents if parent.name == "raw"), None)
            if raw_root is not None:
                root = raw_root.parent
                fallback = resolve_original(root, fallback_match.group(1), target.markdown)
                if fallback is not None:
                    result = extract_html(fallback, title, profile)
                    result.engine = f"fallback-official-html:{result.engine}"
                    result.source_pages = 0
                    return result
    if suffix in {".html", ".htm"}:
        return extract_html(target.original, title, profile)
    if suffix == ".docx":
        return extract_docx(target.original, title, profile)
    if suffix == ".pdf":
        return extract_pdf(target.original, title, profile)
    if suffix == ".md":
        return extract_markdown_passthrough(target.original)
    if suffix == ".doc":
        return extract_markdown_passthrough(target.markdown, legacy_doc=True)
    raise ValueError(f"unsupported source format: {suffix or '<none>'}")


def fallback_existing_result(target: Target, old_body: str, result: ExtractionResult) -> ExtractionResult:
    title = markdown_title(target.metadata, target.markdown)
    original_old = old_body.replace("\r\n", "\n").replace("\r", "\n").strip()
    cleaned_old = original_old
    first_heading = re.match(r"^#\s+(.+?)\n+", original_old)
    if first_heading and comparable_text(first_heading.group(1)) == comparable_text(title):
        cleaned_old = cleaned_old[first_heading.end() :].lstrip()

    if target.original is not None and target.original.suffix.lower() == ".pdf" and result.source_pages:
        asset_dir = f"_assets/{target.markdown.stem}"
        parts = [f"# {title}"]
        for page_number in range(1, result.source_pages + 1):
            parts.extend(
                [
                    f"<!-- source-page: {page_number} -->",
                    f"![Source page {page_number}]({asset_dir}/page-{page_number:03d}.png)",
                ]
            )
        parts.extend(["## Unmapped searchable transcript", cleaned_old])
        headings = [
            Heading(1, title, None, "frontmatter-title", "frontmatter.title"),
            Heading(2, "Unmapped searchable transcript", None, "fallback-marker", "generated-fallback"),
        ]
        return ExtractionResult(
            "\n\n".join(parts).strip() + "\n",
            "fallback-existing-derived-unmapped-pdf",
            result.source_pages,
            headings,
            [],
            sorted(set(result.warnings + ["fallback-existing-derived-unmapped-pdf"])),
            list(range(1, result.source_pages + 1)),
            len(comparable_text(cleaned_old)),
            len(comparable_text(cleaned_old)),
        )

    cleaned_old = original_old
    headings: list[Heading] = []
    for line_number, line in enumerate(cleaned_old.splitlines(), 1):
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if match:
            headings.append(Heading(len(match.group(1)), match.group(2), None, "existing-derived", f"markdown/line[{line_number}]"))
    return ExtractionResult(
        cleaned_old + "\n",
        "fallback-existing-derived-content-regression",
        0,
        headings,
        [],
        sorted(set(result.warnings + ["fallback-existing-derived-content-regression"])),
        [],
        len(comparable_text(cleaned_old)),
        len(comparable_text(cleaned_old)),
    )


def render_pdf_review_pages(target: Target, result: ExtractionResult) -> None:
    if target.original is None or target.original.suffix.lower() != ".pdf":
        return
    if result.engine != "fallback-existing-derived-unmapped-pdf":
        return
    import fitz

    output_dir = target.markdown.parent / "_assets" / target.markdown.stem
    output_dir.mkdir(parents=True, exist_ok=True)
    document = fitz.open(str(target.original))
    try:
        matrix = fitz.Matrix(1.5, 1.5)
        for page_number in result.review_pages:
            destination = output_dir / f"page-{page_number:03d}.png"
            document[page_number - 1].get_pixmap(matrix=matrix, alpha=False).save(str(destination))
    finally:
        document.close()


def structure_path(markdown: Path) -> Path:
    return markdown.with_suffix(".structure.json")


def structure_payload(root: Path, target: Target, result: ExtractionResult, profile: str) -> dict[str, Any]:
    return {
        "schema": "cpa-zh-semantic-structure-v1",
        "generated_at": utc_now(),
        "source_file": relative(root, target.original) if target.original else None,
        "derived_file": relative(root, target.markdown),
        "representation": "semantic-transcript",
        "extraction_profile": profile,
        "structure_status": result.status,
        "extraction_engine": result.engine,
        "source_pages": result.source_pages,
        "heading_count": len(result.headings),
        "table_count": len(result.tables),
        "text_coverage": round(result.coverage, 4),
        "headings": [heading.__dict__ for heading in result.headings],
        "tables": [table.__dict__ for table in result.tables],
        "warnings": result.warnings,
        "review_pages": result.review_pages,
    }


def archive_derived(root: Path, markdown: Path) -> list[Path]:
    archive_root = root / "raw" / "_archive" / "_superseded-derived" / date.today().isoformat()
    archived: list[Path] = []
    for source in (markdown, structure_path(markdown)):
        if not source.exists():
            continue
        rel = source.relative_to(root / "raw")
        destination = archive_root / rel
        shortened = False
        if len(str(destination)) >= 230:
            digest = hashlib.sha256(rel.as_posix().encode("utf-8")).hexdigest()
            destination = archive_root / "_long-paths" / f"{digest}{source.suffix}"
            shortened = True
        destination.parent.mkdir(parents=True, exist_ok=True)
        if destination.exists():
            digest = hashlib.sha256(source.read_bytes()).hexdigest()[:8]
            destination = destination.with_name(f"{destination.stem}-{digest}{destination.suffix}")
        shutil.copy2(source, destination)
        if shortened:
            index_path = archive_root / "archive-index.jsonl"
            with index_path.open("a", encoding="utf-8", newline="\n") as handle:
                handle.write(
                    json.dumps(
                        {"source": relative(root, source), "archived": relative(root, destination)},
                        ensure_ascii=False,
                    )
                    + "\n"
                )
        archived.append(destination)
    return archived


def body_metrics(body: str) -> dict[str, int]:
    return {
        "chars": len(comparable_text(body)),
        "headings": len(markdown_heading_records(body)),
        "tables": body.count("\n<table>") + len(re.findall(r"^\|.*\|$", body, re.MULTILINE)) // 3,
        "page_markers": len(PAGE_MARKER_RE.findall(body)),
    }


def write_reextract_report(path: Path, scope: str, profile: str, apply: bool, rows: list[dict[str, Any]], errors: list[dict[str, str]]) -> None:
    totals = Counter(row["status"] for row in rows)
    lines = [
        "# CPA-ZH Markdown structure re-extraction report",
        "",
        f"- Generated: `{utc_now()}`",
        f"- Scope: `{scope}`",
        f"- Profile: `{profile}`",
        f"- Mode: `{'apply' if apply else 'dry-run'}`",
        f"- Processed: `{len(rows)}`",
        f"- Verified auto: `{totals['verified-auto']}`",
        f"- Needs review: `{totals['needs-review']}`",
        f"- Errors/skipped: `{len(errors)}`",
        "",
        "## Comparison",
        "",
        "| Derived file | Status | Old/New chars | Old/New headings | Old/New tables | Pages | Coverage |",
        "|---|---:|---:|---:|---:|---:|---:|",
    ]
    for row in rows:
        lines.append(
            f"| `{row['path']}` | {row['status']} | {row['old']['chars']}/{row['new']['chars']} | "
            f"{row['old']['headings']}/{row['new']['headings']} | {row['old']['tables']}/{row['new']['tables']} | "
            f"{row['pages']} | {row['coverage']:.1%} |"
        )
    if errors:
        lines.extend(["", "## Errors and skipped files", "", "| File | Reason |", "|---|---|"])
        for item in errors:
            reason = item["reason"].replace("|", "\\|")
            lines.append(f"| `{item['path']}` | {reason} |")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines) + "\n", encoding="utf-8", newline="\n")


def reextract(root: Path, scope: str, profile: str, apply: bool, output: Path) -> int:
    rows: list[dict[str, Any]] = []
    errors: list[dict[str, str]] = []
    targets = list(iter_targets(root, scope))
    for index, target in enumerate(targets, 1):
        rel_path = relative(root, target.markdown)
        if target.original is None:
            errors.append({"path": rel_path, "reason": "original source could not be resolved"})
            continue
        if target.original.suffix.lower() not in SUPPORTED_SUFFIXES:
            errors.append({"path": rel_path, "reason": f"unsupported source format: {target.original.suffix.lower()}"})
            continue
        try:
            metadata, raw_frontmatter, old_body = split_markdown(target.markdown)
            result = extract_target(target, profile)
            old_metrics = body_metrics(old_body)
            new_metrics = body_metrics(result.body)
            if (
                not result.engine.startswith("fallback-official-html:")
                and old_metrics["chars"] >= 200
                and new_metrics["chars"] < old_metrics["chars"] * 0.65
            ):
                result = fallback_existing_result(target, old_body, result)
            updates = {
                "representation": "semantic-transcript",
                "extraction_profile": profile,
                "structure_status": result.status,
                "source_pages": result.source_pages,
                "heading_count": len(result.headings),
                "table_count": len(result.tables),
                "extraction_engine": result.engine,
                "extraction_status": "ok" if result.status == "verified-auto" else "needs-review",
                "structure_updated_at": date.today().isoformat(),
            }
            if result.engine == "fallback-existing-derived-legacy-doc":
                updates["fallback_source"] = "existing-derived-legacy-doc"
            elif result.engine.startswith("fallback-official-html:"):
                updates["fallback_source"] = "official-html-equivalent"
            elif result.engine.startswith("fallback-existing-derived"):
                updates["fallback_source"] = result.engine
            new_document = update_frontmatter(raw_frontmatter, updates) + result.body
            new_metrics = body_metrics(result.body)
            rows.append(
                {
                    "path": rel_path,
                    "status": result.status,
                    "old": old_metrics,
                    "new": new_metrics,
                    "pages": result.source_pages,
                    "coverage": result.coverage,
                }
            )
            if apply:
                archive_derived(root, target.markdown)
                render_pdf_review_pages(target, result)
                target.markdown.write_text(new_document, encoding="utf-8", newline="\n")
                if target.category == "authoritative" and profile == "faithful":
                    sidecar = structure_path(target.markdown)
                    sidecar.write_text(
                        json.dumps(structure_payload(root, target, result, profile), ensure_ascii=False, indent=2) + "\n",
                        encoding="utf-8",
                        newline="\n",
                    )
            if index % 25 == 0:
                print(f"processed={index}/{len(targets)}", flush=True)
        except Exception as exc:
            errors.append({"path": rel_path, "reason": f"{type(exc).__name__}: {exc}"})
    write_reextract_report(output, scope, profile, apply, rows, errors)
    print(f"targets={len(targets)}")
    print(f"processed={len(rows)}")
    print(f"needs_review={sum(1 for row in rows if row['status'] == 'needs-review')}")
    print(f"errors_or_skipped={len(errors)}")
    print(f"report={output}")
    return 1 if errors else 0


def heading_issues(body: str) -> list[str]:
    issues: list[str] = []
    heading_records = markdown_heading_records(body)
    headings = [(level, text) for _, level, text in heading_records]
    previous = 0
    h1_seen: Counter[str] = Counter()
    for level, text in headings:
        normalized = comparable_text(text)
        if level == 1:
            h1_seen[normalized] += 1
        if previous and level > previous + 1:
            issues.append(f"heading-level-jump:{previous}->{level}:{text[:40]}")
        previous = level
    for text, count in h1_seen.items():
        if text and count > 1:
            issues.append(f"duplicate-document-title:{count}:{text[:40]}")

    lines = body.splitlines()
    for (left_line, _, left_text), (right_line, _, right_text) in zip(heading_records, heading_records[1:]):
        between = lines[left_line:right_line - 1]
        has_content = any(line.strip() and not line.strip().startswith("<!--") for line in between)
        if not has_content and comparable_text(left_text) == comparable_text(right_text):
            issues.append(f"consecutive-duplicate-heading:{right_text[:40]}")
    return issues


def audit_target(root: Path, target: Target) -> tuple[list[str], list[str], dict[str, Any]]:
    errors: list[str] = []
    warnings: list[str] = []
    metadata, _, body = split_markdown(target.markdown)
    for key in ("representation", "extraction_profile", "structure_status", "source_pages", "heading_count", "table_count"):
        if key not in metadata:
            errors.append(f"missing-metadata:{key}")
    if target.original is None:
        errors.append("missing-original")
        return errors, warnings, {}
    extraction_engine = str(metadata.get("extraction_engine") or "")
    if target.original.suffix.lower() == ".pdf" and "fallback-official-html" not in extraction_engine:
        try:
            import fitz

            document = fitz.open(str(target.original))
            pages = len(document)
            document.close()
            markers = [int(value) for value in PAGE_MARKER_RE.findall(body)]
            if markers != list(range(1, pages + 1)):
                errors.append(f"pdf-page-markers:{len(markers)}/{pages}")
        except Exception as exc:
            errors.append(f"pdf-open:{type(exc).__name__}")
    warnings.extend(heading_issues(body))
    if re.search(r"[A-Za-z]-\n[A-Za-z]", body):
        warnings.append("word-split-at-line-break")
    if metadata.get("structure_status") == "needs-review":
        warnings.append("declared-needs-review")
    if target.category == "authoritative" and str(metadata.get("extraction_profile")) == "faithful":
        sidecar = structure_path(target.markdown)
        if not sidecar.exists():
            errors.append("missing-structure-json")
    return errors, warnings, body_metrics(body)


def write_audit_report(path: Path, scope: str, rows: list[dict[str, Any]]) -> None:
    error_count = sum(len(row["errors"]) for row in rows)
    warning_count = sum(len(row["warnings"]) for row in rows)
    lines = [
        "# CPA-ZH raw Markdown structure audit",
        "",
        f"- Generated: `{utc_now()}`",
        f"- Scope: `{scope}`",
        f"- Files: `{len(rows)}`",
        f"- Errors: `{error_count}`",
        f"- Warnings: `{warning_count}`",
        f"- Needs review: `{sum('declared-needs-review' in row['warnings'] for row in rows)}`",
        "",
        "## Findings",
        "",
        "| File | Errors | Warnings |",
        "|---|---|---|",
    ]
    for row in rows:
        if not row["errors"] and not row["warnings"]:
            continue
        lines.append(
            f"| `{row['path']}` | {'<br>'.join(row['errors']) or '-'} | {'<br>'.join(row['warnings']) or '-'} |"
        )
    if all(not row["errors"] and not row["warnings"] for row in rows):
        lines.append("| - | - | - |")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines) + "\n", encoding="utf-8", newline="\n")


def audit(root: Path, scope: str, output: Path, write_maintenance_report: bool = False) -> int:
    rows: list[dict[str, Any]] = []
    for target in iter_targets(root, scope):
        errors, warnings, metrics = audit_target(root, target)
        rows.append({"path": relative(root, target.markdown), "errors": errors, "warnings": warnings, "metrics": metrics})
    write_audit_report(output, scope, rows)
    maintenance_report = root / "wiki" / "_maintenance" / "raw-structure-review.md"
    if write_maintenance_report and maintenance_report.resolve() != output.resolve():
        write_audit_report(maintenance_report, scope, rows)
    errors = sum(len(row["errors"]) for row in rows)
    warnings = sum(len(row["warnings"]) for row in rows)
    print(f"files={len(rows)}")
    print(f"errors={errors}")
    print(f"warnings={warnings}")
    print(f"report={output}")
    return 1 if errors else 0


def main() -> int:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    parser = argparse.ArgumentParser(description="Audit and rebuild structure-faithful raw Markdown derivatives.")
    parser.add_argument("--root", default="knowledge-base/CPA-ZH", help="Knowledge base root.")
    subparsers = parser.add_subparsers(dest="command", required=True)
    audit_parser = subparsers.add_parser("audit")
    audit_parser.add_argument("--scope", choices=["authoritative", "cases", "lectures", "all"], default="all")
    audit_parser.add_argument("--output", default="workspace/outputs/raw_structure_audit.md")
    audit_parser.add_argument("--write-maintenance-report", action="store_true")
    extract_parser = subparsers.add_parser("reextract")
    extract_parser.add_argument("--scope", choices=["authoritative", "cases", "lectures", "all"], default="all")
    extract_parser.add_argument("--profile", choices=["faithful", "readable"], default="faithful")
    extract_parser.add_argument("--apply", action="store_true", help="Write files; default is dry-run.")
    extract_parser.add_argument("--output", default="workspace/outputs/raw_structure_reextract_report.md")
    args = parser.parse_args()
    root = Path(args.root).resolve()
    output = Path(args.output)
    if not output.is_absolute():
        output = Path.cwd() / output
    if args.command == "audit":
        return audit(root, args.scope, output, args.write_maintenance_report)
    return reextract(root, args.scope, args.profile, args.apply, output)


if __name__ == "__main__":
    raise SystemExit(main())
