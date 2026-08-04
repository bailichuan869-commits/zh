from __future__ import annotations

import argparse
import csv
import hashlib
import json
import mimetypes
import re
import shutil
import sys
import time
import urllib.error
import urllib.request
from dataclasses import dataclass, asdict
from datetime import date
from pathlib import Path
from urllib.parse import urljoin, urlparse

from bs4 import BeautifulSoup


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_KB = ROOT / "knowledge-base" / "CPA-ZH"
USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 Chrome/126.0 Safari/537.36 CPA-ZH-archiver/1.0"
)
ERROR_MARKERS = ("页面不存在", "不存在或已删除", "您访问的页面", "链接已失效")
ATTACHMENT_SUFFIXES = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".zip"}
REGISTRY_COLUMNS = [
    "article_id", "title", "source_role", "parent_url", "attachment_url",
    "original_path", "markdown_path", "sha256", "content_type",
    "extraction_engine", "extraction_status", "classification", "note",
]


@dataclass
class RegistryItem:
    article_id: str
    title: str
    source_role: str
    parent_url: str
    attachment_url: str
    original_path: str
    markdown_path: str
    sha256: str
    content_type: str
    extraction_engine: str
    extraction_status: str
    classification: str
    note: str = ""


def article_id(url: str) -> str:
    match = re.search(r"t\d+_(\d+)\.htm", url or "")
    return match.group(1) if match else ""


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def normalize_title(value: str) -> str:
    return re.sub(r"[\s—－_（）()《》“”\"'，,。？?：:；;]+", "", value or "").lower()


def yaml_quote(value: object) -> str:
    return json.dumps(str(value), ensure_ascii=False)


def split_frontmatter(text: str) -> tuple[dict[str, str], str]:
    match = re.match(r"^---\n(.*?)\n---\n?(.*)$", text, re.S)
    if not match:
        return {}, text
    meta: dict[str, str] = {}
    for line in match.group(1).splitlines():
        if ":" in line:
            key, value = line.split(":", 1)
            meta[key.strip()] = value.strip().strip('"')
    return meta, match.group(2)


def render_markdown(meta: dict[str, object], body: str) -> str:
    order = [
        "title", "type", "source_type", "source_role", "source_url",
        "attachment_url", "parent_source", "article_id", "original_file",
        "created", "retrieved_at", "converted_at", "sha256",
        "extraction_engine", "extraction_status", "conversion_note",
    ]
    lines = ["---"]
    for key in order:
        value = meta.get(key)
        if value not in (None, ""):
            lines.append(f"{key}: {yaml_quote(value)}")
    for key, value in meta.items():
        if key not in order and value not in (None, ""):
            lines.append(f"{key}: {yaml_quote(value)}")
    lines.extend(["---", "", body.strip(), ""])
    return "\n".join(lines)


def markdown_table(rows: list[dict[str, str]], columns: list[str]) -> str:
    def clean(value: object) -> str:
        return str(value or "").replace("|", "\\|").replace("\n", " ").strip()

    lines = ["| " + " | ".join(columns) + " |", "| " + " | ".join("---" for _ in columns) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(clean(row.get(column, "")) for column in columns) + " |")
    return "\n".join(lines)


def request_bytes(url: str, retries: int = 3) -> tuple[bytes, str]:
    last: Exception | None = None
    for attempt in range(retries):
        try:
            req = urllib.request.Request(url, headers={"User-Agent": USER_AGENT, "Accept": "*/*"})
            with urllib.request.urlopen(req, timeout=45) as response:
                data = response.read()
                content_type = response.headers.get_content_type() or "application/octet-stream"
                return data, content_type
        except (urllib.error.URLError, urllib.error.HTTPError, TimeoutError, OSError) as exc:
            last = exc
            if attempt + 1 < retries:
                time.sleep(1.5 * (attempt + 1))
    raise RuntimeError(f"download failed: {url}: {last}")


def decode_html(data: bytes) -> str:
    for encoding in ("utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def validate_html(data: bytes, url: str, expected_title: str = "") -> str:
    text = decode_html(data)
    if any(marker in text for marker in ERROR_MARKERS):
        raise ValueError(f"official error page returned for {url}")
    soup = BeautifulSoup(text, "lxml")
    title = " ".join((soup.title.get_text(" ", strip=True) if soup.title else "").split())
    if not title:
        title = " ".join((soup.find("h2").get_text(" ", strip=True) if soup.find("h2") else "").split())
    if expected_title:
        left = normalize_title(expected_title)[:24]
        right = normalize_title(title)
        if left and left not in right and right[:20] not in normalize_title(expected_title):
            raise ValueError(f"title mismatch for {url}: expected={expected_title!r}, got={title!r}")
    return text


def validate_attachment(data: bytes, url: str, content_type: str) -> str:
    suffix = Path(urlparse(url).path).suffix.lower()
    if suffix == ".pdf" and not data.startswith(b"%PDF"):
        raise ValueError(f"non-PDF response saved as PDF: {url} ({content_type})")
    if suffix == ".docx" and not data.startswith(b"PK"):
        raise ValueError(f"non-DOCX response: {url} ({content_type})")
    if suffix == ".doc" and data[:8] != bytes.fromhex("D0CF11E0A1B11AE1"):
        raise ValueError(f"non-DOC response: {url} ({content_type})")
    if len(data) < 128:
        raise ValueError(f"attachment too small: {url} ({len(data)} bytes)")
    return suffix


def html_links(html_text: str, base_url: str) -> list[str]:
    soup = BeautifulSoup(html_text, "lxml")
    found: list[str] = []
    for node in soup.select("a[href]"):
        href = str(node.get("href") or "").strip()
        full = urljoin(base_url, href)
        if Path(urlparse(full).path).suffix.lower() in ATTACHMENT_SUFFIXES:
            found.append(full)
    return list(dict.fromkeys(found))


def html_body(html_text: str, title: str, parser: str = "lxml") -> str:
    soup = BeautifulSoup(html_text, parser)
    for selector in ("script", "style", "noscript", ".sharebox", ".noprint", "#footer", ".footer"):
        for node in soup.select(selector):
            node.decompose()
    container = None
    for selector in (".TRS_Editor", ".my_conboxzw", "#UCAP-CONTENT", ".pages_content", "#zoom", ".content", "article"):
        candidate = soup.select_one(selector)
        if candidate and len(candidate.get_text("", strip=True)) > 10:
            container = candidate
            break
    container = container or soup.body or soup
    blocks = container.select("h1, h2, h3, h4, h5, h6, p, li, tr")
    raw_lines = []
    if blocks:
        for block in blocks:
            text = " ".join(block.get_text("", strip=True).split())
            if text:
                raw_lines.append(text)
    else:
        raw_lines = [" ".join(line.split()) for line in container.get_text("\n").splitlines()]
    junk = {
        "热门检索：", "财政收支", "积极财政政策", "减税降费", "首页", "职能机构",
        "新闻报道", "信息公开", "政务服务", "交流互动", "专题专栏", "返回主站",
        "【打印此页】", "【关闭窗口】",
    }
    lines: list[str] = []
    for line in raw_lines:
        if not line or line in junk or line.startswith(("当前位置：", "发布日期：", "发布日期:")):
            continue
        if any(marker in line for marker in ERROR_MARKERS):
            continue
        if line == title and not lines:
            continue
        lines.append(line)
    body = "\n\n".join(lines)
    body = re.sub(r"\n{3,}", "\n\n", body).strip()
    return f"# {title}\n\n{body}".strip()


def pdf_body(data: bytes, title: str) -> tuple[str, str, str]:
    import fitz

    with fitz.open(stream=data, filetype="pdf") as document:
        pages = [page.get_text("text", sort=True).strip() for page in document]
    text = "\n\n".join(part for part in pages if part)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    status = "ok" if len(re.sub(r"\s+", "", text)) >= 80 else "ocr-needed"
    return f"# {title}\n\n{text}".strip(), "pymupdf", status


def docx_body(data: bytes, title: str) -> tuple[str, str, str]:
    import io
    import docx

    document = docx.Document(io.BytesIO(data))
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    text = "\n\n".join(lines)
    return f"# {title}\n\n{text}".strip(), "python-docx", "ok" if len(text) >= 80 else "low-content"


def safe_archive_move(source: Path, destination: Path, apply: bool) -> None:
    if not source.exists():
        return
    if not apply:
        print(f"  MOVE {source} -> {destination}")
        return
    destination.parent.mkdir(parents=True, exist_ok=True)
    if destination.exists():
        if sha256_file(source) == sha256_file(destination):
            source.unlink()
            return
        destination = destination.with_name(destination.stem + "-" + sha256_file(source)[:8] + destination.suffix)
    shutil.move(str(source), str(destination))


class RawRepair:
    def __init__(self, kb: Path, apply: bool) -> None:
        self.kb = kb.resolve()
        self.raw = self.kb / "raw"
        self.archive = self.raw / "_archive"
        self.apply = apply
        self.registry: list[RegistryItem] = []
        self.errors: list[str] = []

    def load_existing_registry(self) -> None:
        path = self.archive / "indexes" / "attachment-registry.csv"
        if not path.exists():
            return
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            for row in csv.DictReader(handle):
                values = {column: str(row.get(column) or "") for column in REGISTRY_COLUMNS}
                if values["attachment_url"]:
                    self.registry.append(RegistryItem(**values))

    def kb_rel(self, path: Path) -> str:
        return path.resolve().relative_to(self.kb).as_posix()

    def write_bytes(self, path: Path, data: bytes) -> None:
        print(f"  WRITE {path} ({len(data)} bytes)")
        if self.apply:
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_bytes(data)

    def write_text(self, path: Path, text: str) -> None:
        print(f"  WRITE {path}")
        if self.apply:
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_text(text, encoding="utf-8", newline="\n")

    def archive_csv(self, name: str) -> Path:
        active = self.raw / "standards" / "accounting" / name
        archived = self.archive / "standards" / "accounting" / name
        return active if active.exists() else archived

    def read_csv(self, name: str) -> list[dict[str, str]]:
        with self.archive_csv(name).open("r", encoding="utf-8-sig", newline="") as fh:
            return list(csv.DictReader(fh))

    def category_articles(self, index_dir: Path) -> dict[str, str]:
        mapping: dict[str, str] = {}
        for path in sorted(index_dir.glob("*.html")):
            slug = path.name.split("-", 1)[0]
            text = path.read_text(encoding="utf-8", errors="ignore")
            for match in re.finditer(r"t\d+_(\d+)\.htm", text):
                key = match.group(1)
                if key in mapping and mapping[key] != slug:
                    raise ValueError(f"article {key} appears in multiple categories")
                mapping[key] = slug
        return mapping

    def corrected_url(self, url: str, slug: str, section: str) -> str:
        marker = f"/{section}/"
        if f"/{section}/{slug}/" in url:
            return url
        return url.replace(marker, f"/{section}/{slug}/", 1)

    def canonical_rows(self, rows: list[dict[str, str]]) -> tuple[list[dict[str, str]], dict[str, list[dict[str, str]]]]:
        grouped: dict[str, list[dict[str, str]]] = {}
        order: list[str] = []
        for row in rows:
            key = article_id(row.get("Url", ""))
            if not key:
                continue
            if key not in grouped:
                grouped[key] = []
                order.append(key)
            grouped[key].append(row)
        return [grouped[key][0] for key in order], grouped

    def migrate_old(self, row: dict[str, str], destination_root: Path) -> None:
        local = Path(row["LocalFile"])
        try:
            relative = local.relative_to(self.raw)
        except ValueError:
            relative = Path("standards/accounting") / local.name
        archived = self.archive / relative
        sidecar = self.raw / Path(str(relative) + ".md")
        safe_archive_move(archived, destination_root / relative, self.apply)
        safe_archive_move(sidecar, destination_root / Path(str(relative) + ".md"), self.apply)

    def canonical_collection_relative(self, kind: str, seq: int, row: dict[str, str]) -> Path:
        if kind == "implementation-qa":
            return Path("standards/accounting/implementation-qa-pages-v2") / f"qa-{seq:03d}.html"
        if kind == "application-case":
            name = Path(row.get("LocalFile", "")).name
            if not name or Path(name).suffix.lower() not in {".html", ".htm"}:
                raise ValueError(f"application-case row has no usable local filename: {row!r}")
            return Path("standards/accounting/application-cases-pages") / name
        raise ValueError(f"unsupported accounting collection kind: {kind}")

    def archive_accidental_nested_outputs(self) -> None:
        destination = self.archive / "_invalid-runs" / "2026-07-28-path-resolution-bug"
        nested = self.archive / "_archive" / "standards" / "accounting"
        if nested.exists():
            for path in sorted(item for item in nested.rglob("*") if item.is_file()):
                relative = path.relative_to(nested)
                safe_archive_move(path, destination / "nested-archive" / relative, self.apply)

        for directory_name in ("implementation-qa-pages-v2", "application-cases-pages"):
            directory = self.archive / "standards" / "accounting" / directory_name
            if not directory.exists():
                continue
            for path in sorted(directory.glob("*.md")):
                relative = path.relative_to(self.archive / "standards" / "accounting")
                safe_archive_move(path, destination / "archived-sidecars" / relative, self.apply)

    def archive_superseded_legacy_qas(self) -> None:
        relative_root = Path("standards/accounting/implementation-qa-pages")
        original_root = self.archive / relative_root
        markdown_root = self.raw / relative_root
        destination = self.archive / "_superseded" / relative_root
        if original_root.exists():
            for path in sorted(item for item in original_root.rglob("*") if item.is_file()):
                safe_archive_move(path, destination / path.relative_to(original_root), self.apply)
        if markdown_root.exists():
            for path in sorted(item for item in markdown_root.rglob("*") if item.is_file()):
                safe_archive_move(path, destination / path.relative_to(markdown_root), self.apply)

    def normalize_known_source_roles(self) -> None:
        policy_root = self.raw / "policies" / "second-section"
        if policy_root.exists():
            for path in sorted(policy_root.glob("*/official.html.md")):
                has_attachment = any(path.parent.glob("*.pdf.md")) or any((path.parent / "attachments").glob("*.pdf.md"))
                if has_attachment:
                    self.set_source_role(path, "attachment-landing")
        self.set_source_role(
            self.raw / "standards/accounting/mof-accounting-standards-topic.html.md",
            "index-page",
        )

        audit_page_urls = {
            "cicpa-guidelines-34-20230410.html.md": "https://www.cicpa.org.cn/xxfb/tzgg/202304/t20230410_64066.html",
            "cicpa-guidelines-15-20220120.html.md": "https://www.cicpa.org.cn/xxfb/tzgg/202201/t20220120_63335.html",
            "cicpa-standards-20220120.html.md": "https://www.cicpa.org.cn/xxfb/tzgg/202201/t20220120_63336.html",
            "cicpa-standards-20230103.html.md": "https://www.cicpa.org.cn/xxfb/tzgg/202301/t20230103_63902.html",
        }
        audit_root = self.raw / "standards" / "audit"
        for filename, source_url in audit_page_urls.items():
            self.set_source_metadata(audit_root / filename, source_url=source_url, source_role="content")

    def refresh_full_text_policy_pages(self) -> None:
        slugs = (
            "audit-order-2021-30",
            "caihui-supervision-2023-4",
            "cpa-exam-2024-115",
            "firm-license-supervision-2019-97",
        )
        for slug in slugs:
            directory = self.raw / "policies" / "second-section" / slug
            original = directory / "official.html"
            if not original.exists():
                archived = self.archive / "policies" / "second-section" / slug / "official.html"
                original = archived if archived.exists() else original
            markdown = directory / "official.html.md"
            source_file = directory / "source-url.txt"
            if not original.exists() or not markdown.exists() or not source_file.exists():
                continue
            old_meta, _old_body = split_frontmatter(markdown.read_text(encoding="utf-8", errors="ignore"))
            title = old_meta.get("title") or slug
            source_url = source_file.read_text(encoding="utf-8", errors="ignore").strip()
            data = original.read_bytes()
            html_text = validate_html(data, source_url, title)
            body = html_body(html_text, title, parser="html.parser")
            if len(re.sub(r"\s+", "", body)) < 500:
                raise ValueError(f"policy full text too short after extraction: {source_url}")
            self.write_source_md(original, markdown, title, source_url, "content", body, data)

    def set_source_role(self, path: Path, role: str) -> None:
        if not path.exists():
            return
        text = path.read_text(encoding="utf-8", errors="ignore")
        meta, body = split_frontmatter(text)
        if meta.get("source_role") == role:
            return
        meta["source_role"] = role
        self.write_text(path, render_markdown(meta, body))

    def set_source_metadata(self, path: Path, *, source_url: str, source_role: str) -> None:
        if not path.exists():
            return
        text = path.read_text(encoding="utf-8", errors="ignore")
        meta, body = split_frontmatter(text)
        changed = False
        values = {
            "source_url": source_url,
            "source_role": source_role,
            "article_id": article_id(source_url),
        }
        for key, value in values.items():
            if meta.get(key) != value:
                meta[key] = value
                changed = True
        if changed:
            self.write_text(path, render_markdown(meta, body))

    @staticmethod
    def is_markdown_table_separator(line: str) -> bool:
        stripped = line.strip()
        if not stripped.startswith("|") or not stripped.endswith("|"):
            return False
        cells = [cell.strip() for cell in stripped[1:-1].split("|")]
        return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)

    def normalize_markdown_tables(self) -> None:
        for path in sorted(self.raw.rglob("*.md")):
            if "_archive" in path.parts:
                continue
            text = path.read_text(encoding="utf-8", errors="ignore")
            lines = text.splitlines()
            cleaned: list[str] = []
            changed = False
            for line in lines:
                if (
                    cleaned
                    and self.is_markdown_table_separator(line)
                    and self.is_markdown_table_separator(cleaned[-1])
                ):
                    changed = True
                    continue
                cleaned.append(line)
            if changed:
                self.write_text(path, "\n".join(cleaned) + ("\n" if text.endswith("\n") else ""))

    def refresh_enriched_markdown_manifest(self) -> None:
        manifest = self.raw / "lectures/ai-coding-lectures-2026-07-09/manifest.json"
        if not manifest.exists():
            return
        try:
            data = json.loads(manifest.read_text(encoding="utf-8"))
        except Exception as exc:
            self.errors.append(f"cannot read lecture manifest {manifest}: {exc}")
            return
        if not isinstance(data, list):
            return
        changed = False
        for item in data:
            if not isinstance(item, dict):
                continue
            local_file = str(item.get("local_file") or "").replace("\\", "/")
            if not local_file:
                continue
            local_path = self.kb / local_file
            if not local_path.exists() or local_path.suffix.lower() != ".md":
                continue
            values = {"bytes": local_path.stat().st_size, "sha256": sha256_file(local_path)}
            for key, value in values.items():
                if item.get(key) != value:
                    item[key] = value
                    changed = True
        if changed:
            self.write_text(manifest, json.dumps(data, ensure_ascii=False, indent=2) + "\n")

    def write_source_md(self, archived: Path, active_md: Path, title: str, url: str, role: str, body: str, data: bytes, engine: str = "beautifulsoup", status: str = "ok", attachment_url: str = "", equivalent_source_url: str = "", conversion_note: str = "") -> None:
        suffix = archived.suffix.lower()
        if suffix == ".pdf":
            source_type = "pdf"
        elif suffix in {".doc", ".docx"}:
            source_type = "doc"
        elif suffix in {".xls", ".xlsx", ".csv"}:
            source_type = "table"
        else:
            source_type = "web-snapshot"
        meta = {
            "title": title,
            "type": "raw-source",
            "source_type": source_type,
            "source_role": role,
            "source_url": url,
            "attachment_url": attachment_url,
            "equivalent_source_url": equivalent_source_url,
            "article_id": article_id(url),
            "original_file": self.kb_rel(archived),
            "created": date.today().isoformat(),
            "retrieved_at": date.today().isoformat(),
            "converted_at": date.today().isoformat(),
            "sha256": sha256_bytes(data),
            "extraction_engine": engine,
            "extraction_status": status,
            "conversion_note": conversion_note,
        }
        self.write_text(active_md, render_markdown(meta, body))

    def official_html_fallback(self, attachment_url: str, title: str) -> tuple[str, str, str] | None:
        fallbacks = {
            "https://kjs.mof.gov.cn/zhengcefabu/201211/P020121122516762746185.doc": (
                "https://kjs.mof.gov.cn/zhengcefabu/201211/t20121122_701308.htm",
                Path("standards/accounting/interpretations-pages/关于印发企业会计准则解释第5号的通知.html"),
                "七、本解释自2013年1月1日施行",
            ),
            "https://kjs.mof.gov.cn/zhengcefabu/201309/P020230310389475958559.doc": (
                "https://www.gov.cn/gongbao/content/2013/content_2528121.htm",
                Path("standards/accounting/equivalent-html-fallbacks/enterprise-product-cost-accounting-system-trial.html"),
                "第五十三条",
            ),
            "https://kjs.mof.gov.cn/zhengcefabu/201612/P020161212572651599508.pdf": (
                "https://kjs.mof.gov.cn/zhengcefabu/201612/t20161212_2479869.htm",
                Path("standards/accounting/other-rules-pages/009-关于印发-增值税会计处理规定-的通知.html"),
                "四、附则",
            ),
        }
        fallback = fallbacks.get(attachment_url)
        if fallback is None:
            return None
        source_url, relative, required_marker = fallback
        original = self.archive / relative
        markdown = self.raw / Path(str(relative) + ".md")
        if original.exists():
            data = original.read_bytes()
        else:
            data, _content_type = request_bytes(source_url)
            self.write_bytes(original, data)
        html_text = validate_html(data, source_url, title)
        parser = "html.parser" if source_url == "https://www.gov.cn/gongbao/content/2013/content_2528121.htm" else "lxml"
        body = html_body(html_text, title, parser=parser)
        if required_marker not in body:
            raise ValueError(f"official HTML fallback is incomplete ({required_marker!r} missing): {source_url}")
        if not markdown.exists():
            self.write_source_md(
                original, markdown, title, source_url, "equivalent-source", body, data,
                engine="beautifulsoup", status="ok",
            )
        return body, source_url, self.kb_rel(original)

    def register_attachment(self, *, title: str, parent_url: str, attachment_url: str, original: Path, markdown: Path, data: bytes, content_type: str, engine: str, status: str, classification: str = "archived", note: str = "") -> None:
        self.registry.append(RegistryItem(
            article_id=article_id(parent_url), title=title, source_role="substantive-attachment",
            parent_url=parent_url, attachment_url=attachment_url,
            original_path=self.kb_rel(original), markdown_path=self.kb_rel(markdown),
            sha256=sha256_bytes(data), content_type=content_type,
            extraction_engine=engine, extraction_status=status,
            classification=classification, note=note,
        ))

    def download_attachment(self, title: str, parent_url: str, attachment_url: str, original: Path, markdown: Path) -> None:
        try:
            if original.exists() and markdown.exists():
                data = original.read_bytes()
                content_type = mimetypes.guess_type(original.name)[0] or "application/octet-stream"
            else:
                data, content_type = request_bytes(attachment_url)
            suffix = validate_attachment(data, attachment_url, content_type)
            fallback = self.official_html_fallback(attachment_url, title)
            if fallback is not None:
                body, equivalent_source_url, equivalent_original = fallback
                engine, status = "official-html-fallback", "fallback-official-html"
                classification = "equivalent-source-archived"
                note = f"Binary original retained; text derived from {equivalent_original} ({equivalent_source_url})"
            else:
                equivalent_source_url = ""
                classification = "archived"
                note = ""
                if suffix == ".pdf":
                    body, engine, status = pdf_body(data, title)
                elif suffix == ".docx":
                    body, engine, status = docx_body(data, title)
                else:
                    body = f"# {title}\n\n该原件为 `{suffix}` 格式，当前环境未直接抽取正文。"
                    engine, status = "binary-archive", "fallback-required"
            if not (original.exists() and markdown.exists()):
                self.write_bytes(original, data)
            self.write_source_md(
                original, markdown, title, parent_url, "substantive-attachment", body, data,
                engine, status, attachment_url, equivalent_source_url,
                note if fallback is not None else "",
            )
            self.register_attachment(
                title=title, parent_url=parent_url, attachment_url=attachment_url,
                original=original, markdown=markdown, data=data, content_type=content_type,
                engine=engine, status=status, classification=classification, note=note,
            )
        except Exception as exc:
            message = f"{attachment_url}: {exc}"
            self.errors.append(message)
            print(f"  ERROR {message}")

    def repair_accounting_collection(self, *, csv_name: str, index_dir: str, section: str, expected: int, kind: str) -> None:
        rows = self.read_csv(csv_name)
        canonical, grouped = self.canonical_rows(rows)
        slug_map = self.category_articles(self.archive / "standards" / "accounting" / index_dir)
        if len(canonical) != expected or set(slug_map) != set(grouped):
            raise ValueError(f"{kind}: expected {expected}, csv={len(canonical)}, indexes={len(slug_map)}")
        output_rows: list[dict[str, str]] = []
        invalid_root = self.archive / "_invalid-downloads"
        duplicate_root = self.archive / "_duplicates"

        for seq, row in enumerate(canonical, 1):
            key = article_id(row["Url"])
            url = self.corrected_url(row["Url"], slug_map[key], section)
            title = row["Title"].strip()
            relative = self.canonical_collection_relative(kind, seq, row)
            archived = self.archive / relative
            active_md = self.raw / Path(str(relative) + ".md")
            try:
                replacing = True
                if archived.exists():
                    data = archived.read_bytes()
                    try:
                        html_text = validate_html(data, url, title)
                        replacing = False
                    except ValueError:
                        data, _content_type = request_bytes(url)
                        html_text = validate_html(data, url, title)
                else:
                    data, _content_type = request_bytes(url)
                    html_text = validate_html(data, url, title)
                role = "attachment-landing" if kind == "application-case" else "content"
                body = html_body(html_text, title)
                if kind == "implementation-qa" and len(re.sub(r"\s+", "", body)) < 100:
                    raise ValueError(f"implementation answer too short: {url}")
                # Preserve the bad snapshot and duplicates only after the replacement is validated.
                if replacing:
                    self.migrate_old(row, invalid_root)
                    for duplicate in grouped[key][1:]:
                        self.migrate_old(duplicate, duplicate_root)
                    self.write_bytes(archived, data)
                self.write_source_md(archived, active_md, title, url, role, body, data)
                attachment_urls = html_links(html_text, url)
                if kind == "application-case":
                    pdf_urls = [item for item in attachment_urls if Path(urlparse(item).path).suffix.lower() == ".pdf"]
                    if len(pdf_urls) != 1:
                        raise ValueError(f"expected one application-case PDF, got {len(pdf_urls)}: {url}")
                    attachment_name = f"{seq:03d}-{Path(urlparse(pdf_urls[0]).path).name}"
                    attachment_original = self.archive / "standards" / "accounting" / "application-case-attachments" / attachment_name
                    attachment_md = self.raw / "standards" / "accounting" / "application-case-attachments" / f"{attachment_name}.md"
                    self.download_attachment(title, url, pdf_urls[0], attachment_original, attachment_md)
                updated = dict(row)
                updated.update({
                    "Seq": str(seq), "Url": url, "LocalFile": str(archived),
                    "DerivedMarkdown": str(active_md), "Status": "ok",
                    "ArticleId": key, "CategorySlug": slug_map[key],
                })
                output_rows.append(updated)
            except Exception as exc:
                message = f"{kind} {seq}/{key}: {exc}"
                self.errors.append(message)
                print(f"  ERROR {message}")

        if len(output_rows) == expected:
            columns = list(dict.fromkeys(key for row in output_rows for key in row))
            csv_text = self.csv_text(output_rows, columns)
            csv_path = self.archive / "standards" / "accounting" / csv_name
            self.write_text(csv_path, csv_text)
            md_path = self.raw / "standards" / "accounting" / f"{csv_name}.md"
            meta = {
                "title": Path(csv_name).stem, "type": "raw-source", "source_type": "table",
                "source_role": "source-registry", "original_file": self.kb_rel(csv_path),
                "created": date.today().isoformat(), "converted_at": date.today().isoformat(),
                "extraction_engine": "csv", "extraction_status": "ok",
            }
            self.write_text(md_path, render_markdown(meta, f"# {Path(csv_name).stem}\n\n" + markdown_table(output_rows, columns)))

    @staticmethod
    def csv_text(rows: list[dict[str, str]], columns: list[str]) -> str:
        import io
        output = io.StringIO(newline="")
        writer = csv.DictWriter(output, fieldnames=columns, extrasaction="ignore", lineterminator="\n")
        writer.writeheader()
        writer.writerows(rows)
        return output.getvalue()

    def add_existing_audit_registry(self) -> None:
        csv_path = self.archive / "standards" / "audit" / "downloaded-cicpa-professional-standards-pdfs.csv"
        if not csv_path.exists():
            return
        with csv_path.open("r", encoding="utf-8-sig", newline="") as fh:
            rows = list(csv.DictReader(fh))
        for row in rows:
            local = Path(row.get("LocalFile", ""))
            try:
                relative = local.relative_to(self.raw)
            except ValueError:
                continue
            original = self.archive / relative
            markdown = self.raw / Path(str(relative) + ".md")
            if original.exists() and markdown.exists():
                self.registry.append(RegistryItem(
                    article_id="", title=row.get("Title", ""), source_role="substantive-attachment",
                    parent_url="", attachment_url=row.get("Url", ""),
                    original_path=self.kb_rel(original), markdown_path=self.kb_rel(markdown),
                    sha256=sha256_file(original), content_type="application/pdf",
                    extraction_engine="existing-sidecar", extraction_status="ok",
                    classification="archived", note="CICPA direct-PDF CSV mapping",
                ))

    def add_manifest_registry(self) -> None:
        for manifest in self.raw.rglob("manifest.json"):
            if "_archive" in manifest.parts:
                continue
            try:
                data = json.loads(manifest.read_text(encoding="utf-8"))
            except Exception:
                continue
            items = data if isinstance(data, list) else data.get("items", [])
            if not isinstance(items, list):
                continue
            for item in items:
                if not isinstance(item, dict):
                    continue
                url = str(item.get("attachment_url") or "")
                local_rel = str(item.get("local_file") or "").replace("\\", "/")
                md_rel = str(item.get("derived_markdown") or "").replace("\\", "/")
                if not url or not local_rel or not md_rel:
                    continue
                original = self.kb / local_rel
                if not original.exists() and local_rel.startswith("raw/"):
                    original = self.archive / Path(local_rel).relative_to("raw")
                markdown = self.kb / md_rel
                if original.exists() and markdown.exists():
                    self.registry.append(RegistryItem(
                        article_id="", title=str(item.get("title") or item.get("slug") or original.stem),
                        source_role="substantive-attachment", parent_url=str(item.get("url") or ""),
                        attachment_url=url, original_path=self.kb_rel(original), markdown_path=self.kb_rel(markdown),
                        sha256=sha256_file(original), content_type=str(item.get("content_type") or mimetypes.guess_type(original.name)[0] or "application/octet-stream"),
                        extraction_engine="manifest", extraction_status=str(item.get("text_extraction_status") or "ok"),
                        classification="archived", note="manifest mapping",
                    ))

    def register_independence_equivalent_attachment(self) -> None:
        mof_url = "http://kjs.mof.gov.cn/zhengcefabu/202501/P020250120381094517836.pdf"
        cicpa_url = "https://cicpa.org.cn/xxfb/news/202501/W020250120543364207300.pdf"
        original = self.archive / "ethics/third-section/independence-standard-2024-29-pdf/official.pdf"
        markdown = self.raw / "ethics/third-section/independence-standard-2024-29-pdf/official.pdf.md"
        if not original.exists() or not markdown.exists():
            return
        self.registry.append(RegistryItem(
            article_id="3952051",
            title="中国注册会计师独立性准则第1号",
            source_role="substantive-attachment",
            parent_url="http://kjs.mof.gov.cn/zhengcefabu/202501/t20250120_3952051.htm",
            attachment_url=mof_url,
            original_path=self.kb_rel(original),
            markdown_path=self.kb_rel(markdown),
            sha256=sha256_file(original),
            content_type="application/pdf",
            extraction_engine="existing-sidecar",
            extraction_status="ok",
            classification="equivalent-source-archived",
            note=f"Equivalent official CICPA copy archived from {cicpa_url}",
        ))

    def register_existing_audit_bundle(self) -> None:
        attachment_url = "https://www.cicpa.org.cn/xxfb/tzgg/202301/W020231121556491370245.zip"
        original = self.raw / "standards/audit/archives/2023-23-audit-standards.zip"
        extracted = self.raw / "standards/audit/archives/2023-23-audit-standards"
        if not original.exists() or not extracted.exists():
            return
        extracted_count = sum(1 for path in extracted.rglob("*") if path.is_file())
        self.registry.append(RegistryItem(
            article_id="63902",
            title="中国注册会计师审计准则第1101号等23项准则",
            source_role="substantive-attachment",
            parent_url="https://www.cicpa.org.cn/xxfb/tzgg/202301/t20230103_63902.html",
            attachment_url=attachment_url,
            original_path=self.kb_rel(original),
            markdown_path=self.kb_rel(extracted),
            sha256=sha256_file(original),
            content_type="application/zip",
            extraction_engine="zipfile",
            extraction_status="ok",
            classification="archived",
            note=f"Aggregate ZIP retained and expanded to {extracted_count} files",
        ))

    def repair_named_attachments(self, *, download_missing: bool = True) -> None:
        targets = [
            (
                "会计师事务所监督检查办法", "https://www.gov.cn/zhengce/zhengceku/2022-05/16/content_5690682.htm",
                "https://www.gov.cn/zhengce/zhengceku/2022-05/16/5690682/files/6573ca570b0245f6ab3e0a7adf14324c.pdf",
                "policies/second-section/firm-inspection-2022-23/attachments/official.pdf",
            ),
            (
                "注册会计师行业诚信建设纲要", "https://www.gov.cn/zhengce/zhengceku/2023-04/02/content_5749779.htm",
                "https://www.gov.cn/zhengce/zhengceku/2023-04/02/5749779/files/784a06d3e09545b28aafb171b137adfe.pdf",
                "policies/second-section/integrity-2023-5/attachments/official.pdf",
            ),
            (
                "独立性准则第1号应用指南征求意见稿起草说明", "https://www.cicpa.org.cn/xxfb/tzgg/202504/t20250430_65411.html",
                "https://www.cicpa.org.cn/xxfb/tzgg/202504/W020250430561278344331.pdf",
                "ethics/third-section/independence-application-guide-exposure-2025/attachments/drafting-notes.pdf",
            ),
        ]
        for title, parent_url, attachment_url, relative in targets:
            original = self.archive / relative
            markdown = self.raw / f"{relative}.md"
            if original.exists() and markdown.exists():
                data = original.read_bytes()
                self.register_attachment(
                    title=title, parent_url=parent_url, attachment_url=attachment_url,
                    original=original, markdown=markdown, data=data,
                    content_type="application/pdf", engine="existing-sidecar", status="ok",
                    note="existing named attachment",
                )
                self.set_source_metadata(markdown, source_url=parent_url, source_role="substantive-attachment")
            elif download_missing:
                self.download_attachment(title, parent_url, attachment_url, original, markdown)
        self.registry.append(RegistryItem(
            article_id="", title="意见反馈表", source_role="auxiliary-attachment",
            parent_url="https://www.cicpa.org.cn/xxfb/tzgg/202504/t20250430_65411.html",
            attachment_url="https://www.cicpa.org.cn/xxfb/tzgg/202504/W020250430561278352757.docx",
            original_path="", markdown_path="", sha256="", content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            extraction_engine="excluded", extraction_status="not-required", classification="excluded-auxiliary",
            note="征求意见反馈表，不属于规范正文",
        ))

    def repair_accounting_page_attachments(self) -> None:
        collections = [
            "downloaded-enterprise-accounting-standards-interpretations.csv",
            "downloaded-enterprise-accounting-standards-other-rules.csv",
        ]
        known_urls = {item.attachment_url for item in self.registry if item.classification == "archived"}
        for csv_name in collections:
            for seq, row in enumerate(self.read_csv(csv_name), 1):
                parent_url = row.get("Url", "")
                local = Path(row.get("LocalFile", ""))
                try:
                    relative = local.relative_to(self.raw)
                except ValueError:
                    continue
                archived_page = self.archive / relative
                if not archived_page.exists():
                    continue
                html_text = archived_page.read_text(encoding="utf-8", errors="ignore")
                for attachment_url in html_links(html_text, parent_url):
                    if attachment_url in known_urls:
                        continue
                    suffix = Path(urlparse(attachment_url).path).suffix.lower()
                    title = row.get("Title", "")
                    interpretation_match = re.search(r"解释第\s*(\d+)\s*号", title)
                    existing_original = None
                    existing_md = None
                    if interpretation_match:
                        no = int(interpretation_match.group(1))
                        candidates = list((self.raw / "standards/accounting/interpretations-pages").glob(f"解释第{no:02d}号_附件.*"))
                        candidates = [path for path in candidates if not path.name.endswith(".md")]
                        if candidates:
                            existing_original = candidates[0]
                            existing_md = Path(str(existing_original) + ".md")
                    if existing_original and existing_md and existing_md.exists():
                        data = existing_original.read_bytes()
                        self.register_attachment(
                            title=title, parent_url=parent_url, attachment_url=attachment_url,
                            original=existing_original, markdown=existing_md, data=data,
                            content_type=mimetypes.guess_type(existing_original.name)[0] or "application/octet-stream",
                            engine="existing-sidecar", status="ok", note="renamed interpretation attachment",
                        )
                        known_urls.add(attachment_url)
                        continue
                    basename = Path(urlparse(attachment_url).path).name
                    stem = f"{seq:03d}-{basename}"
                    original = self.archive / "standards/accounting/page-attachments" / stem
                    markdown = self.raw / "standards/accounting/page-attachments" / f"{stem}.md"
                    self.download_attachment(title, parent_url, attachment_url, original, markdown)
                    known_urls.add(attachment_url)

    def enrich_frontmatter(self) -> None:
        role_by_name = re.compile(r"(?:index|official-links|registry|README)", re.I)
        for path in sorted(self.raw.rglob("*.md")):
            if "_archive" in path.parts:
                continue
            text = path.read_text(encoding="utf-8", errors="ignore")
            meta, body = split_frontmatter(text)
            changed = False
            if "source_type" not in meta:
                original = meta.get("original_file", "")
                suffix = Path(original).suffix.lower()
                if suffix in {".html", ".htm", ".xml"}:
                    source_type = "web-snapshot"
                elif suffix == ".pdf":
                    source_type = "pdf"
                elif suffix in {".doc", ".docx"}:
                    source_type = "doc"
                elif suffix == ".csv":
                    source_type = "table"
                elif "lectures" in path.parts:
                    source_type = "local-lecture"
                elif "laws" in path.parts:
                    source_type = "law"
                else:
                    source_type = "generated-index" if role_by_name.search(path.name) else "local-source"
                meta["source_type"] = source_type
                changed = True
            if "source_role" not in meta:
                if meta.get("source_type") == "table":
                    role = "source-registry"
                else:
                    role = "index-page" if role_by_name.search(path.name) or "indexes" in path.parts else "content"
                meta["source_role"] = role
                changed = True
            if changed:
                self.write_text(path, render_markdown(meta, body))

    def write_registry(self) -> None:
        unique: dict[tuple[str, str], RegistryItem] = {}
        for item in self.registry:
            key = (item.attachment_url, item.original_path)
            current = unique.get(key)
            if (
                current is None
                or current.classification == item.classification
                or (current.classification != "archived" and item.classification == "archived")
            ):
                unique[key] = item
        rows = [asdict(item) for item in sorted(unique.values(), key=lambda row: (row.classification, row.attachment_url, row.original_path))]
        csv_path = self.archive / "indexes" / "attachment-registry.csv"
        self.write_text(csv_path, self.csv_text(rows, REGISTRY_COLUMNS))
        md_path = self.raw / "indexes" / "attachment-registry.csv.md"
        meta = {
            "title": "CPA-ZH 附件登记表", "type": "raw-source", "source_type": "table",
            "source_role": "source-registry", "original_file": self.kb_rel(csv_path),
            "created": date.today().isoformat(), "converted_at": date.today().isoformat(),
            "extraction_engine": "csv", "extraction_status": "ok",
        }
        self.write_text(md_path, render_markdown(meta, "# CPA-ZH 附件登记表\n\n" + markdown_table(rows, REGISTRY_COLUMNS)))

    def run(self, scope: str) -> int:
        self.load_existing_registry()
        if scope in {"accounting", "all"}:
            self.archive_accidental_nested_outputs()
            self.archive_superseded_legacy_qas()
            self.repair_accounting_collection(
                csv_name="downloaded-enterprise-accounting-standards-implementation-qa-v2.csv",
                index_dir="implementation-qa-indexes", section="sswd", expected=100,
                kind="implementation-qa",
            )
            self.repair_accounting_collection(
                csv_name="downloaded-enterprise-accounting-standards-application-cases.csv",
                index_dir="application-cases-indexes", section="srzzzq", expected=39,
                kind="application-case",
            )
        self.add_existing_audit_registry()
        self.add_manifest_registry()
        self.register_independence_equivalent_attachment()
        self.register_existing_audit_bundle()
        self.repair_named_attachments(download_missing=scope in {"policies", "ethics", "all"})
        if scope in {"accounting", "all"}:
            self.repair_accounting_page_attachments()
        if scope in {"policies", "all"}:
            self.refresh_full_text_policy_pages()
        self.enrich_frontmatter()
        self.normalize_known_source_roles()
        self.normalize_markdown_tables()
        self.refresh_enriched_markdown_manifest()
        self.write_registry()
        print(f"registry_items={len(self.registry)} errors={len(self.errors)} apply={self.apply}")
        for error in self.errors:
            print(f"ERROR: {error}")
        return 1 if self.errors else 0


def main() -> int:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    parser = argparse.ArgumentParser(description="Repair and register CPA-ZH raw official sources.")
    parser.add_argument("--root", default=str(DEFAULT_KB), help="Knowledge base root.")
    parser.add_argument("--scope", choices=["accounting", "policies", "ethics", "all"], default="all")
    parser.add_argument("--apply", action="store_true", help="Write changes. Default is dry-run.")
    args = parser.parse_args()
    return RawRepair(Path(args.root), args.apply).run(args.scope)


if __name__ == "__main__":
    raise SystemExit(main())
