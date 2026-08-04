#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
convert_raw_to_md.py — 把 raw/ 下各种格式的原文件统一转换为 Markdown。

策略（用户已拍板：内容提取式 + 原文件归档）：
  1. 对每个可转换文件（html/htm/xml/pdf/docx/txt/csv）抽取正文内容，
     生成一个同名的 .md 门面文件（统一 frontmatter + 提取的正文）。
  2. 原文件移动到 raw/_archive/<相对路径> 归档（目录保持原结构）。
  3. 可选 --update-wiki：把 wiki 概念页里 `raw/.../xxx.EXT` 引用更新为
     `raw/.../xxx.EXT.md`，使 UI 跳转到新的 md 而不是已归档的原文件。
  4. JSON（manifest/metadata 索引数据）不转正文，仅作 frontmatter 补充来源。

frontmatter 字段：
  title / type / source_type / original_file / source_url / document_no
  / created / converted_at / tags

用法：
  python tools/convert_raw_to_md.py --dry-run            # 仅打印计划
  python tools/convert_raw_to_md.py --dir raw/standards/accounting --apply   # 试点一个目录
  python tools/convert_raw_to_md.py --apply --update-wiki                 # 全量
  python tools/convert_raw_to_md.py --clean            # 后处理清洗已转换 md（dry-run）
  python tools/convert_raw_to_md.py --clean --apply    # 实际清洗：剥导航垃圾 + 补 source_type
"""
import argparse
import csv
import json
import re
import shutil
import sys
from datetime import date
from pathlib import Path

ROOT = Path("knowledge-base/CPA-ZH")
RAW = ROOT / "raw"
WIKI = ROOT / "wiki"
ARCHIVE = RAW / "_archive"

CONVERTIBLE = {".html", ".htm", ".xml", ".pdf", ".docx", ".txt", ".csv"}
SKIP_NAMES = {"metadata.json", "manifest.json", "source-url.txt"}
SKIP_DIRS = {"_archive", "cache", ".git"}

# ---------------------------------------------------------------------------
# 导航/模板垃圾指纹：财政部网站通用模板（返回主站/财政部微信/页脚/面包屑等）
# ---------------------------------------------------------------------------
NAV_CLASS_HINTS = {
    "nav", "navbar", "navigation", "header", "topbar", "top-bar", "footer",
    "foot", "menu", "menubar", "sidebar", "side-bar", "breadcrumb", "crumb",
    "share", "sharebox", "wechat", "weixin", "wx", "qr", "ewm", "ewcode",
    "site-info", "siteinfo", "copyright", "gover", "gov", "banner", "toolbar",
    "loginbar", "mfwx",
    # mof.gov.cn 通用模板：顶部 logo + 部门名 + 站内搜索 + 当前位置 + 返回主站 + 右下二维码
    "logodiv", "zzname", "zz_serach", "searchinput", "buttoninput",
    "returnmain", "dangqian", "popfr",
}

def _build_nav_res():
    exact = [
        r"^返回主站$", r"^返回首页$", r"^财政部微信$", r"^网站地图$",
        r"^主办单位：中华人民共和国财政部$",
        r"^网站标识码：bm\d+$", r"^京ICP备\d+号$", r"^京公网安备\d+号$",
        r"^技术支持：财政部信息网络中心$", r"^首页>$",
    ]
    contains = [
        r"版权所有", r"ICP备[\s\d-]*号?", r"网站标识码：?bm\d*", r"备案号",
        r"京公网安备[\s\d-]*号?", r"无障碍浏览?", r"记住本站", r"收藏本站",
        r"扫一扫", r"微信扫一扫", r"二维码",
        r"当前位置[:：].*?(?=。|$)", r"您当前的位置", r"您所在位置",
        r"分享至朋友圈", r"【大中小】",
    ]
    return [re.compile(p) for p in exact], [re.compile(p) for p in contains]

EXACT_NAV_RES, CONTAINS_NAV_RES = _build_nav_res()
STRONG_NAV_RE = re.compile(r"返回主站|财政部微信|网站地图|主办单位：中华人民共和国财政部")
MEGAMENU_RE = re.compile(r"^[\u4e00-\u9fa5]{2,6}(?:\s*/\s*[\u4e00-\u9fa5]{2,6}){2,}$")


def decompose_nav(soup):
    """按 CSS class/id 命中导航/页脚/侧栏/分享等块并删除，避免模板泄漏。"""
    for el in soup.find_all(True):
        am = el.attrs
        if not am:
            continue
        cls = " ".join(am.get("class", []) or [])
        idv = am.get("id", "") or ""
        attrs = (cls + " " + idv).lower()
        tag = el.name
        if tag in {"nav", "header", "footer", "aside"} or any(h in attrs for h in NAV_CLASS_HINTS):
            el.decompose()


def strip_boilerplate(text: str) -> str:
    """抽取后过滤模板行：
    - 精确整行垃圾（返回主站/财政部微信/网站地图/页脚块等）直接丢弃；
    - 含模板子串的短行（纯模板）整行丢弃，长行（正文被尾随模板污染，如页脚被
      merge_fragments 融合进正文末行）仅剥离子串，保留正文；
    - 仅当文件已被判定为'污染'时，才额外丢弃巨菜单行（避免误删正文
      '收入/成本/利润' 这类合法行）。"""
    lines = text.split("\n")
    polluted = any(STRONG_NAV_RE.search(ln) for ln in lines)
    out = []
    for ln in lines:
        s = ln.strip()
        if not s:
            if out and out[-1] != "":
                out.append("")
            continue
        # 1) 精确整行垃圾 → 丢弃
        if any(r.search(s) for r in EXACT_NAV_RES):
            continue
        # 2) 含模板子串 → 短行丢整行，长行仅剥离子串
        drop = False
        changed = False
        for r in CONTAINS_NAV_RES:
            if r.search(s):
                if len(s) <= 30:
                    drop = True
                    break
                s = r.sub("", s).strip()
                changed = True
        if drop:
            continue
        if changed:
            s = s.strip(" 、，。,.;；:：·")
        # 3) 污染文件的巨菜单行
        if polluted and MEGAMENU_RE.match(s):
            continue
        out.append(s if changed else ln)
    merged = []
    for ln in out:
        if ln == "" and merged and merged[-1] == "":
            continue
        merged.append(ln)
    return "\n".join(merged).strip()

# ---------------------------------------------------------------------------
# YAML 安全转义（只处理冒号/引号/换行，够用）
# ---------------------------------------------------------------------------
def yq(v: str) -> str:
    s = str(v).replace("\\", "\\\\").replace('"', '\\"')
    if any(ch in s for ch in (":", "#", "-", "[", "]", "{", "}", "\n")) or s.strip() != s:
        s = '"' + s + '"'
    return s

# ---------------------------------------------------------------------------
# 正文抽取
# ---------------------------------------------------------------------------
def extract_html(path: Path) -> str:
    from bs4 import BeautifulSoup
    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "nav", "header", "footer", "aside", "form", "iframe", "noscript"]):
        tag.decompose()
    decompose_nav(soup)
    main = None
    for sel in ["div.TRS_Editor", "div.my_doccontent", "div.my_conboxzw", "#zoom",
                "div.box_content",
                "article", "div.article", "div.content", "div.main", "main", "div.v_news_content"]:
        el = soup.select_one(sel)
        if el and len(el.get_text(strip=True)) > 10:
            main = el
            break
    if main is None:
        main = soup.body or soup
    text = main.get_text("\n", strip=True)
    text = re.sub(r"[ \t]+", " ", text)
    text = merge_fragments(text)
    text = reparagraph(text)
    text = strip_boilerplate(text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def merge_fragments(text: str) -> str:
    """合并被 <br>/span 打断的中文碎片行，让正文读起来连贯。"""
    out: list[str] = []
    for ln in text.split("\n"):
        ln = ln.strip()
        if not ln:
            if out and out[-1] != "":
                out.append("")
            continue
        if out and out[-1] != "":
            prev = out[-1]
            if len(ln) <= 3 or re.match(r"^第.+条$", prev) or re.match(r"^第[一二三四五六七八九十]+章", prev):
                out[-1] = prev + ln
                continue
        out.append(ln)
    return "\n".join(out)


def reparagraph(text: str) -> str:
    """按章/条重新分段，贴近法规条文的标准排版。"""
    text = re.sub(r"第([一二三四五六七八九十百零]+)章", r"\n\n第\1章", text)
    text = re.sub(r"第([0-9]{1,3}|[一二三四五六七八九十百零]+|[XIV]+)条", r"\n第\1条", text)
    return text.strip()

def extract_pdf(path: Path) -> str:
    from PyPDF2 import PdfReader
    try:
        r = PdfReader(str(path))
        parts = [(pg.extract_text() or "") for pg in r.pages]
        text = "\n\n".join(p.strip() for p in parts if p.strip())
        return re.sub(r"\n{3,}", "\n\n", text).strip()
    except Exception as e:
        return f"（PDF 解析失败：{e}）"

def extract_docx(path: Path) -> str:
    import docx
    d = docx.Document(str(path))
    out = []
    for p in d.paragraphs:
        if p.text.strip():
            out.append(p.text.strip())
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            out.append(" | ".join(cells))
    return "\n\n".join(out).strip()

def extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore").strip()

def extract_csv(path: Path, limit: int = 80) -> str:
    rows = list(csv.reader(path.read_text(encoding="utf-8", errors="ignore").splitlines()))
    if not rows:
        return ""
    rows = rows[:limit]
    lines = ["| " + " | ".join(c.replace("\n", " ").strip() for c in r) + " |" for r in rows]
    # 表头下加分隔行
    if len(lines) > 1:
        sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
        lines.insert(1, sep)
    total = len(rows)
    head = f"_（共 {total} 行，此处展示前 {min(limit, total)} 行）_\n\n" if total > limit else ""
    return head + "\n".join(lines)

# ---------------------------------------------------------------------------
# frontmatter 来源：同目录 metadata.json / source-url.txt
# ---------------------------------------------------------------------------
def sidecar_meta(path: Path) -> dict:
    meta = {}
    d = path.parent
    mj = d / "metadata.json"
    if mj.exists():
        try:
            data = json.loads(mj.read_text(encoding="utf-8"))
            meta["title"] = data.get("title")
            meta["source_url"] = data.get("url") or data.get("source_url")
            meta["document_no"] = data.get("document_no")
            meta["created"] = data.get("published") or data.get("date") or data.get("created")
            if data.get("tags"):
                meta["tags"] = ",".join(data["tags"]) if isinstance(data["tags"], list) else str(data["tags"])
        except Exception:
            pass
    su = d / "source-url.txt"
    if su.exists() and not meta.get("source_url"):
        meta["source_url"] = su.read_text(encoding="utf-8").strip().splitlines()[0] if su.read_text(encoding="utf-8").strip() else ""
    return {k: v for k, v in meta.items() if v}

def clean_title(stem: str) -> str:
    # 去前缀编号 058- / 123、 等；下划线转空格；保留原文件名中的连字符
    s = re.sub(r"^[0-9]{1,4}[-_．、\s]+", "", stem)
    s = s.replace("_", " ")
    return s.strip() or stem

TYPE_MAP = {
    ".html": "web-snapshot", ".htm": "web-snapshot", ".xml": "web-snapshot",
    ".pdf": "pdf", ".docx": "doc", ".txt": "text", ".csv": "table",
}

def build_md(src: Path) -> tuple[str, dict]:
    ext = src.suffix.lower()
    if ext == ".html" or ext == ".htm" or ext == ".xml":
        body = extract_html(src)
    elif ext == ".pdf":
        body = extract_pdf(src)
    elif ext == ".docx":
        body = extract_docx(src)
    elif ext == ".txt":
        body = extract_txt(src)
    elif ext == ".csv":
        body = extract_csv(src)
    else:
        body = ""
    if not body or len(body) < 5:
        body = f"（无法从原文件提取可识别文本，原类型 {ext}，请查看归档原文 `_archive/{src.relative_to(RAW)}`）"
    meta = sidecar_meta(src)
    meta["title"] = meta.get("title") or clean_title(src.stem)
    meta["type"] = "raw-source"
    meta["source_type"] = TYPE_MAP.get(ext, "other")
    meta["original_file"] = f"raw/{src.relative_to(RAW).as_posix()}"
    meta["converted_at"] = date.today().isoformat()
    if not meta.get("created"):
        meta["created"] = date.fromtimestamp(src.stat().st_mtime).isoformat()
    fm = ["---"]
    for k in ("title", "type", "source_type", "original_file", "source_url",
              "document_no", "created", "converted_at", "tags"):
        v = meta.get(k)
        if v:
            fm.append(f"{k}: {yq(v)}")
    fm.append("---")
    return "\n".join(fm) + "\n\n" + body + "\n", meta

# ---------------------------------------------------------------------------
# 已转换 md 的后处理清洗 + frontmatter 工具
# ---------------------------------------------------------------------------
def split_frontmatter(text: str) -> tuple[dict, str]:
    m = re.match(r"^---\n(.*?)\n---\n?(.*)$", text, re.DOTALL)
    if m:
        fm_raw, body = m.group(1), m.group(2)
        meta = {}
        for line in fm_raw.splitlines():
            if ":" in line:
                k, v = line.split(":", 1)
                meta[k.strip()] = v.strip().strip('"')
        return meta, body
    return {}, text

def dump_frontmatter(meta: dict, body: str) -> str:
    fm = ["---"]
    ordered = ("title", "type", "source_type", "source_role", "original_file", "source_url",
               "attachment_url", "parent_source", "article_id", "document_no", "created",
               "retrieved_at", "converted_at", "sha256", "extraction_engine",
               "extraction_status", "conversion_note", "tags")
    for k in ordered:
        v = meta.get(k)
        if v:
            fm.append(f"{k}: {yq(v)}")
    for k, v in meta.items():
        if k not in ordered and v:
            fm.append(f"{k}: {yq(v)}")
    fm.append("---")
    return "\n".join(fm) + "\n\n" + body + "\n"

def _norm_body(s: str) -> str:
    """仅比较非空行序列，忽略空行数量与首尾空白，用于判定是否真有内容行被删。"""
    return "\n".join(ln.strip() for ln in s.split("\n") if ln.strip())


def clean_existing_md(subdir, apply):
    """后处理清洗已转换的 raw/*.md：剥导航垃圾 + 补缺失 source_type（幂等，仅改动确有变化的文件）。"""
    if subdir:
        sub = subdir.lstrip("/")
        if sub.startswith("raw/"):
            sub = sub[4:]
        base = RAW / sub
    else:
        base = RAW
    if not base.exists():
        print(f"[clean] 目录不存在: {base}")
        return
    targets = [p for p in sorted(base.rglob("*.md")) if p.is_file()]
    n_clean = 0
    n_boiler = 0
    n_meta = 0
    for p in targets:
        rel = p.relative_to(RAW)
        if any(part in SKIP_DIRS for part in rel.parts):
            continue
        if "_archive" in rel.parts or "_trash" in rel.parts or "_maintenance" in rel.parts:
            continue
        t = p.read_text(encoding="utf-8")
        meta, body = split_frontmatter(t)
        new_body = strip_boilerplate(body)
        boiler_changed = _norm_body(body) != _norm_body(new_body)
        meta_missing = "source_type" not in meta and bool(meta.get("original_file"))
        if not (boiler_changed or meta_missing):
            continue
        n_clean += 1
        if boiler_changed:
            n_boiler += 1
        if meta_missing:
            meta["source_type"] = TYPE_MAP.get(Path(meta["original_file"]).suffix.lower(), "other")
            n_meta += 1
        out_body = new_body if boiler_changed else body
        if apply:
            p.write_text(dump_frontmatter(meta, out_body), encoding="utf-8")
        print(f"  {'CLEAN' if apply else 'DRY'}  {rel}  [boiler={boiler_changed}, meta={meta_missing}]")
    print(f"[clean] 待处理 {n_clean} 个（剥导航垃圾 {n_boiler}，补 source_type {n_meta}），{'已写入' if apply else 'dry-run 未改动'}")

# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------
def iter_targets(subdir: str | None):
    base = RAW / subdir if subdir else RAW
    for p in sorted(base.rglob("*")):
        if not p.is_file():
            continue
        if p.name in SKIP_NAMES:
            continue
        if any(part in SKIP_DIRS for part in p.relative_to(RAW).parts):
            continue
        if p.suffix.lower() in CONVERTIBLE:
            yield p

def convert_all(subdir, apply, update_wiki):
    sub = (subdir or "").lstrip("/")
    if sub.startswith("raw/"):
        sub = sub[4:]
    targets = list(iter_targets(sub or None))
    print(f"[plan] 可转换文件: {len(targets)}" + (f" (限定 {subdir})" if subdir else ""))
    done = 0
    for src in targets:
        rel = src.relative_to(RAW)
        md_path = src.with_suffix(src.suffix + ".md")
        md, meta = build_md(src)
        arc = ARCHIVE / rel
        if not apply:
            tag = " [覆盖]" if md_path.exists() else ""
            print(f"  DRY  {rel}  ->  {rel}.md{tag}  [{meta['source_type']}, {len(md)}B]")
            continue
        # apply：已存在则直接覆盖写（不删除，避免触发安全删除拦截）
        md_path.write_text(md, encoding="utf-8")
        arc.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(src), str(arc))
        done += 1
        if done % 50 == 0:
            print(f"  ... 已转换 {done}/{len(targets)}")
    if apply:
        print(f"[done] 实际转换 {done} 个文件，原文件归档至 {ARCHIVE.relative_to(ROOT)}/")
    if update_wiki and apply:
        update_wiki_refs()

def update_wiki_refs():
    pat = re.compile(r"`(raw/[^\s`\\]+\.(?:html|htm|xml|pdf|docx|txt|csv))`")
    cnt = 0
    for p in sorted(WIKI.rglob("*.md")):
        if any(part in ("_trash", "_maintenance") for part in p.parts):
            continue
        t = p.read_text(encoding="utf-8")
        nt = pat.sub(lambda m: "`" + m.group(1) + ".md`", t)
        if nt != t:
            p.write_text(nt, encoding="utf-8")
            cnt += 1
    print(f"[wiki] 更新了 {cnt} 个 wiki 概念页的 raw 引用 -> .md")

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--dir", help="限定子目录，如 raw/standards/accounting（试点用）")
    ap.add_argument("--apply", action="store_true", help="实际执行（默认 dry-run）")
    ap.add_argument("--update-wiki", action="store_true", help="同时更新 wiki 概念页的 raw 引用")
    ap.add_argument("--clean", action="store_true", help="后处理清洗已转换 raw/*.md（剥导航垃圾 + 补 source_type）")
    args = ap.parse_args()
    if args.clean:
        clean_existing_md(args.dir, args.apply)
        return
    convert_all(args.dir, args.apply, args.update_wiki)

if __name__ == "__main__":
    main()
